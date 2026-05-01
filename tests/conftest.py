"""Shared pytest fixtures for the docx-editor test suite."""
from __future__ import annotations

import os
import sys

import pytest
from docx import Document
from docx.shared import Pt

# Make scripts/ importable
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../scripts")))


# ─── Document factories ───────────────────────────────────────────────────────

def _save(doc: Document, path: str) -> str:
    doc.save(path)
    return path


@pytest.fixture
def blank_doc(tmp_path) -> str:
    """Empty document with no paragraphs."""
    doc = Document()
    # Remove the default empty paragraph python-docx adds
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    return _save(doc, str(tmp_path / "blank.docx"))


@pytest.fixture
def simple_doc(tmp_path) -> str:
    """Document with 5 plain text paragraphs."""
    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.add_paragraph("Third paragraph")
    doc.add_paragraph("Fourth paragraph")
    doc.add_paragraph("Fifth paragraph")
    return _save(doc, str(tmp_path / "simple.docx"))


@pytest.fixture
def heading_doc(tmp_path) -> str:
    """Document with headings and body text."""
    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Intro body text.")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("Methods body text.")
    doc.add_heading("Results", level=2)
    doc.add_paragraph("Results body text.")
    return _save(doc, str(tmp_path / "headings.docx"))


@pytest.fixture
def multirun_doc(tmp_path) -> str:
    """Document where a phrase is split across multiple runs with different formatting."""
    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    # Build a paragraph: "Hello" (bold) + " " (plain) + "World" (italic)
    para = doc.add_paragraph()
    run1 = para.add_run("Hello")
    run1.bold = True
    run2 = para.add_run(" ")
    run3 = para.add_run("World")
    run3.italic = True
    # Second paragraph: plain "foo bar baz"
    doc.add_paragraph("foo bar baz")
    return _save(doc, str(tmp_path / "multirun.docx"))


@pytest.fixture
def table_doc(tmp_path) -> str:
    """Document with a 3x3 table."""
    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    headers = ["Name", "Age", "City"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [("Alice", "30", "New York"), ("Bob", "25", "London")]
    for r, row_data in enumerate(data, start=1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val
    return _save(doc, str(tmp_path / "table.docx"))


@pytest.fixture
def large_doc(tmp_path) -> str:
    """1000-paragraph document for performance tests."""
    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    for i in range(1000):
        if i % 50 == 0:
            doc.add_heading(f"Section {i // 50 + 1}", level=1)
        else:
            doc.add_paragraph(f"Paragraph {i}: The quick brown fox jumps over the lazy dog. " * 3)
    return _save(doc, str(tmp_path / "large.docx"))
