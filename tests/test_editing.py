"""Tests for docx_engine.editing — insert, delete, replace."""
from __future__ import annotations

import os
import sys

import pytest
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../scripts")))

from docx_engine import editing
from docx_engine.core import load_document


# ─── insert_paragraph ────────────────────────────────────────────────────────

class TestInsertParagraph:
    def test_append_to_end(self, simple_doc):
        res = editing.insert_paragraph(simple_doc, "New last paragraph")
        assert "SUCCESS" in res
        doc, _ = load_document(simple_doc)
        assert doc.paragraphs[-1].text == "New last paragraph"

    def test_insert_at_index(self, simple_doc):
        res = editing.insert_paragraph(simple_doc, "Inserted", index=0)
        assert "SUCCESS" in res
        doc, _ = load_document(simple_doc)
        # Inserted AFTER index 0, so it's at position 1
        assert doc.paragraphs[1].text == "Inserted"

    def test_insert_invalid_index(self, simple_doc):
        res = editing.insert_paragraph(simple_doc, "X", index=999)
        assert res.startswith("ERROR")

    def test_insert_negative_index(self, simple_doc):
        res = editing.insert_paragraph(simple_doc, "X", index=-1)
        assert res.startswith("ERROR")

    def test_insert_after_heading(self, heading_doc):
        res = editing.insert_paragraph(heading_doc, "After intro", after_heading="Introduction")
        assert "SUCCESS" in res
        doc, _ = load_document(heading_doc)
        texts = [p.text for p in doc.paragraphs]
        intro_idx = texts.index("Introduction")
        assert texts[intro_idx + 1] == "After intro"

    def test_insert_after_heading_not_found(self, simple_doc):
        res = editing.insert_paragraph(simple_doc, "X", after_heading="Nonexistent")
        assert res.startswith("ERROR")

    def test_insert_with_style(self, simple_doc):
        res = editing.insert_paragraph(simple_doc, "Styled", style="Normal")
        assert "SUCCESS" in res


# ─── delete_paragraphs ───────────────────────────────────────────────────────

class TestDeleteParagraphs:
    def test_delete_single(self, simple_doc):
        doc, _ = load_document(simple_doc)
        original_count = len(doc.paragraphs)
        res = editing.delete_paragraphs(simple_doc, [1])
        assert "SUCCESS" in res
        doc, _ = load_document(simple_doc)
        assert len(doc.paragraphs) == original_count - 1

    def test_delete_multiple(self, simple_doc):
        doc, _ = load_document(simple_doc)
        original_count = len(doc.paragraphs)
        res = editing.delete_paragraphs(simple_doc, [0, 2])
        assert "SUCCESS" in res
        doc, _ = load_document(simple_doc)
        assert len(doc.paragraphs) == original_count - 2

    def test_delete_string_indices(self, simple_doc):
        res = editing.delete_paragraphs(simple_doc, "0,1")
        assert "SUCCESS" in res

    def test_delete_out_of_range(self, simple_doc):
        res = editing.delete_paragraphs(simple_doc, [999])
        assert res.startswith("ERROR")

    def test_delete_preserves_content(self, simple_doc):
        res = editing.delete_paragraphs(simple_doc, [1])
        assert "SUCCESS" in res
        doc, _ = load_document(simple_doc)
        texts = [p.text for p in doc.paragraphs]
        assert "First paragraph" in texts
        assert "Second paragraph" not in texts
        assert "Third paragraph" in texts


# ─── replace_text ────────────────────────────────────────────────────────────

class TestReplaceText:
    def test_simple_replace(self, simple_doc):
        res = editing.replace_text(simple_doc, "First", "1st")
        assert "SUCCESS" in res
        doc, _ = load_document(simple_doc)
        assert doc.paragraphs[0].text == "1st paragraph"

    def test_replace_not_found(self, simple_doc):
        res = editing.replace_text(simple_doc, "XXXXNOTFOUND", "Y")
        assert "not found" in res.lower() or res.startswith("'XXXX")

    def test_replace_regex(self, simple_doc):
        res = editing.replace_text(simple_doc, r"\w+th", "Nth", use_regex=True)
        assert "SUCCESS" in res

    def test_replace_counts_occurrences(self, simple_doc):
        # "paragraph" appears in all 5 paragraphs
        res = editing.replace_text(simple_doc, "paragraph", "section")
        assert "5" in res

    def test_replace_preserves_other_paragraphs(self, simple_doc):
        editing.replace_text(simple_doc, "First", "ONE")
        doc, _ = load_document(simple_doc)
        assert doc.paragraphs[1].text == "Second paragraph"
        assert doc.paragraphs[2].text == "Third paragraph"


# ─── _surgical_replace (multi-run correctness) ───────────────────────────────

class TestSurgicalReplace:
    """Directly test the surgical replace helper for multi-run formatting preservation."""

    def _make_para_with_runs(self, texts: list[str]) -> object:
        """Create an in-memory paragraph with one run per text item."""
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        para = doc.add_paragraph()
        for t in texts:
            para.add_run(t)
        return para

    def test_single_run_replace(self):
        para = self._make_para_with_runs(["hello world"])
        count = editing._surgical_replace(para, "hello", "hi")
        assert count == 1
        assert para.runs[0].text == "hi world"

    def test_multi_run_replace_cross_boundary(self):
        # "Hello World" split across 3 runs; replace "o W" which crosses run 0→1→2
        para = self._make_para_with_runs(["Hell", "o ", "World"])
        count = editing._surgical_replace(para, "o W", "0-W")
        assert count == 1
        full = "".join(r.text for r in para.runs)
        assert full == "Hell0-World"

    def test_multi_run_replace_preserves_surrounding(self):
        # Replace text that spans runs; content before and after must survive
        para = self._make_para_with_runs(["abcde", "fghij"])
        count = editing._surgical_replace(para, "defg", "X")
        assert count == 1
        full = "".join(r.text for r in para.runs)
        assert full == "abcXhij"

    def test_replace_at_run_start(self):
        para = self._make_para_with_runs(["abc", "def"])
        count = editing._surgical_replace(para, "abcd", "Z")
        assert count == 1
        full = "".join(r.text for r in para.runs)
        assert full == "Zef"

    def test_replace_spanning_all_runs(self):
        para = self._make_para_with_runs(["a", "b", "c"])
        count = editing._surgical_replace(para, "abc", "X")
        assert count == 1
        full = "".join(r.text for r in para.runs)
        assert full == "X"

    def test_multiple_occurrences(self):
        para = self._make_para_with_runs(["foo bar foo"])
        count = editing._surgical_replace(para, "foo", "baz")
        assert count == 2
        full = "".join(r.text for r in para.runs)
        assert full == "baz bar baz"

    def test_no_match_returns_zero(self):
        para = self._make_para_with_runs(["hello"])
        count = editing._surgical_replace(para, "NOTFOUND", "X")
        assert count == 0
        assert para.runs[0].text == "hello"

    def test_regex_replace(self):
        para = self._make_para_with_runs(["foo123bar"])
        count = editing._surgical_replace(para, r"\d+", "NUM", use_regex=True)
        assert count == 1
        assert para.runs[0].text == "fooNUMbar"

    def test_formatting_preserved_in_untouched_runs(self):
        """Bold/italic on runs not touched by replacement must be unchanged."""
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        para = doc.add_paragraph()
        r1 = para.add_run("Hello ")
        r1.bold = True
        r2 = para.add_run("cruel ")
        r3 = para.add_run("World")
        r3.italic = True

        editing._surgical_replace(para, "cruel ", "")
        assert r1.bold is True
        assert r3.italic is True
