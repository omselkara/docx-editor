"""Edge case and robustness tests — empty docs, unicode, boundary conditions."""
from __future__ import annotations

import os
import sys

import pytest
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../scripts")))

from docx_engine import core, reading, editing, formatting, tables
from docx_engine.core import load_document


class TestMissingFile:
    def test_load_missing_file(self):
        doc, err = load_document("/nonexistent/path/file.docx")
        assert doc is None
        assert err.startswith("ERROR")

    def test_reading_missing_file(self):
        res = reading.get_outline("/nonexistent.docx")
        assert res.startswith("ERROR")

    def test_editing_missing_file(self):
        res = editing.insert_paragraph("/nonexistent.docx", "text")
        assert res.startswith("ERROR")


class TestEmptyDocument:
    def test_outline_empty(self, blank_doc):
        res = reading.get_outline(blank_doc, json_mode=True)
        assert res["status"] == "WARNING"

    def test_full_text_empty(self, blank_doc):
        res = reading.full_text(blank_doc, json_mode=True)
        assert res["status"] == "SUCCESS"
        assert res["paragraphs"] == []

    def test_replace_text_empty(self, blank_doc):
        res = editing.replace_text(blank_doc, "anything", "replacement")
        assert "not found" in res.lower() or "no changes" in res.lower()

    def test_delete_paragraphs_empty(self, blank_doc):
        res = editing.delete_paragraphs(blank_doc, [0])
        assert res.startswith("ERROR")

    def test_read_range_empty(self, blank_doc):
        res = reading.read_range(blank_doc, 0, 0, json_mode=True)
        assert res["status"] == "ERROR"


class TestUnicode:
    def test_unicode_paragraph(self, tmp_path):
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        unicode_texts = [
            "Türkçe: çğışöü",
            "日本語: こんにちは",
            "Arabic: مرحبا",
            "Emoji: 🎉🚀",
            "Mixed: héllo wörld",
        ]
        for text in unicode_texts:
            doc.add_paragraph(text)
        path = str(tmp_path / "unicode.docx")
        doc.save(path)

        res = reading.full_text(path, json_mode=True)
        assert res["status"] == "SUCCESS"
        extracted = [p["text"] for p in res["paragraphs"]]
        for text in unicode_texts:
            assert text in extracted

    def test_unicode_replace(self, tmp_path):
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        doc.add_paragraph("Merhaba dünya")
        path = str(tmp_path / "turkish.docx")
        doc.save(path)

        res = editing.replace_text(path, "dünya", "world")
        assert "SUCCESS" in res
        doc2, _ = load_document(path)
        assert doc2.paragraphs[0].text == "Merhaba world"


class TestBoundaryIndices:
    def test_read_range_last_valid_index(self, simple_doc):
        doc, _ = load_document(simple_doc)
        n = len(doc.paragraphs)
        res = reading.read_range(simple_doc, n - 1, n - 1, json_mode=True)
        assert res["status"] == "SUCCESS"

    def test_read_range_exactly_out_of_bounds(self, simple_doc):
        doc, _ = load_document(simple_doc)
        n = len(doc.paragraphs)
        # end_idx == n is out of range (valid is 0..n-1)
        res = reading.read_range(simple_doc, 0, n, json_mode=True)
        assert res["status"] == "ERROR"

    def test_insert_at_last_index(self, simple_doc):
        doc, _ = load_document(simple_doc)
        last = len(doc.paragraphs) - 1
        res = editing.insert_paragraph(simple_doc, "After last", index=last)
        assert "SUCCESS" in res

    def test_insert_at_zero(self, simple_doc):
        res = editing.insert_paragraph(simple_doc, "After zero", index=0)
        assert "SUCCESS" in res

    def test_format_text_out_of_range_index_skipped(self, simple_doc):
        res = formatting.format_text(simple_doc, para_indices=[9999], bold=True)
        assert res.startswith("WARNING")

    def test_table_index_zero_when_no_tables(self, simple_doc):
        res = tables.read_table(simple_doc, 0, json_mode=True)
        assert res["status"] == "ERROR"


class TestSpecialContent:
    def test_paragraph_with_only_whitespace(self, tmp_path):
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        doc.add_paragraph("   ")
        doc.add_paragraph("real content")
        path = str(tmp_path / "whitespace.docx")
        doc.save(path)

        res = reading.full_text(path, compact=True, json_mode=True)
        assert res["status"] == "SUCCESS"

    def test_very_long_paragraph(self, tmp_path):
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        doc.add_paragraph("A" * 10000)
        path = str(tmp_path / "long.docx")
        doc.save(path)

        res = reading.full_text(path, json_mode=True)
        assert res["paragraphs"][0]["text"] == "A" * 10000

    def test_replace_in_table_cell(self, table_doc):
        """replace_text must also replace inside table cells."""
        res = editing.replace_text(table_doc, "Alice", "Alicia")
        assert "SUCCESS" in res
        doc, _ = load_document(table_doc)
        cell_texts = [cell.text for row in doc.tables[0].rows for cell in row.cells]
        assert "Alicia" in cell_texts
        assert "Alice" not in cell_texts
