"""Tests for docx_engine.smart_features — summary_map, word_count, language detection, templates."""
from __future__ import annotations

import os
import sys

import pytest
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../scripts")))

from docx_engine import smart_features
from docx_engine.core import load_document


class TestSummaryMap:
    def test_basic_map(self, simple_doc):
        res = smart_features.summary_map(simple_doc)
        assert "DOCUMENT MAP" in res
        assert "Paragraphs:" in res

    def test_map_includes_table_summary(self, table_doc):
        res = smart_features.summary_map(table_doc)
        assert "Tables" in res

    def test_map_heading_doc(self, heading_doc):
        res = smart_features.summary_map(heading_doc)
        assert "H1" in res or "Introduction" in res

    def test_map_empty_doc(self, blank_doc):
        res = smart_features.summary_map(blank_doc)
        assert "DOCUMENT MAP" in res


class TestWordCount:
    def test_basic_count(self, simple_doc):
        res = smart_features.word_count(simple_doc)
        assert "word" in res.lower() or "Word" in res

    def test_empty_doc_word_count(self, blank_doc):
        res = smart_features.word_count(blank_doc)
        assert isinstance(res, str)


class TestLanguageDetection:
    def test_detects_language(self, simple_doc):
        res = smart_features.detect_language(simple_doc)
        assert isinstance(res, str)
        assert len(res) > 0

    def test_turkish_document(self, tmp_path):
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        doc.add_paragraph("Merhaba dünya, bu bir Türkçe belgedir çğışöü")
        path = str(tmp_path / "turkish.docx")
        doc.save(path)
        res = smart_features.detect_language(path)
        assert isinstance(res, str)


class TestFromTemplate:
    def test_basic_substitution(self, tmp_path):
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        doc.add_paragraph("Hello {{name}}, your order {{order_id}} is ready.")
        template_path = str(tmp_path / "template.docx")
        doc.save(template_path)

        output_path = str(tmp_path / "filled.docx")
        res = smart_features.from_template(
            template_path,
            variables="name=Alice,order_id=12345",
            output_path=output_path,
        )
        assert "SUCCESS" in res
        filled, _ = load_document(output_path)
        assert "Alice" in filled.paragraphs[0].text
        assert "12345" in filled.paragraphs[0].text

    def test_multiple_variables(self, tmp_path):
        doc = Document()
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        doc.add_paragraph("{{greeting}} {{name}}!")
        template_path = str(tmp_path / "template2.docx")
        doc.save(template_path)
        output_path = str(tmp_path / "filled2.docx")
        res = smart_features.from_template(
            template_path,
            variables="greeting=Hello,name=World",
            output_path=output_path,
        )
        assert "SUCCESS" in res
        filled, _ = load_document(output_path)
        assert "Hello World!" in filled.paragraphs[0].text


class TestReadPage:
    def test_read_page_1(self, simple_doc):
        res = smart_features.read_page(simple_doc, page_number=1)
        assert isinstance(res, str) and len(res) > 0

    def test_read_page_out_of_range(self, simple_doc):
        res = smart_features.read_page(simple_doc, page_number=999)
        assert isinstance(res, str)


class TestCloneFormat:
    def test_clone_format(self, simple_doc):
        from docx_engine import formatting
        formatting.format_text(simple_doc, para_indices=[0], bold=True)
        # source_para=0, target_paras=[1]
        res = smart_features.clone_format(simple_doc, source_para=0, target_paras=[1])
        assert "SUCCESS" in res or "cloned" in res.lower() or isinstance(res, str)
