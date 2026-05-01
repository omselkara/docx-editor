"""Comments, footnotes, endnotes, tracked changes, and bookmarks — full annotation support."""
import datetime
from typing import List, Optional, Dict, Any, Union
from lxml import etree
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx_engine.core import load_document, save_document, NSMAP
from docx_engine import errors


# ===================== COMMENTS =====================

def read_comments(doc_path: str) -> str:
    """Read all comments from the document, including author, date, and referenced text."""
    doc, err = load_document(doc_path)
    if err:
        return err

    # Access the comments part directly from the package
    comments_part = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith('/comments'):
            comments_part = rel.target_part
            break

    if comments_part is None:
        return "No comments found in the document."

    comments_xml = comments_part._element
    comments = comments_xml.findall(qn('w:comment'))

    if not comments:
        return "No comments found in the document."

    # Build a map of comment IDs to their referenced text
    body = doc.element.body
    comment_refs = {}
    for cs in body.iter(qn('w:commentRangeStart')):
        cid = cs.get(qn('w:id'))
        ce = body.find(f'.//{qn("w:commentRangeEnd")}[@{qn("w:id")}="{cid}"]')
        if ce is not None:
            # Collect text between commentRangeStart and commentRangeEnd
            ref_text = _get_text_between(cs, ce)
            comment_refs[cid] = ref_text

    lines = ["=== COMMENTS ==="]
    for comment in comments:
        cid = comment.get(qn('w:id'))
        author = comment.get(qn('w:author'), '(unknown)')
        date = comment.get(qn('w:date'), '(no date)')
        initials = comment.get(qn('w:initials'), '')

        # Get comment text
        comment_text_parts = []
        for p in comment.findall(qn('w:p')):
            runs = p.findall(f'.//{qn("w:t")}')
            text = ''.join(r.text or '' for r in runs)
            comment_text_parts.append(text)
        comment_text = '\n'.join(comment_text_parts)

        ref_text = comment_refs.get(cid, '(reference text not found)')

        lines.append(f"\n  [COMMENT #{cid}]")
        lines.append(f"    Author: {author} ({initials})")
        lines.append(f"    Date: {date}")
        lines.append(f"    Reference text: \"{ref_text[:100]}{'...' if len(ref_text) > 100 else ''}\"")
        lines.append(f"    Comment: {comment_text}")

    return "\n".join(lines)


def _get_text_between(start_elem: Any, end_elem: Any) -> str:
    """Get all text content between two XML elements."""
    text_parts = []
    collecting = False
    for elem in start_elem.getparent().iter():
        if elem is start_elem:
            collecting = True
            continue
        if elem is end_elem:
            break
        if collecting and elem.tag == qn('w:t'):
            text_parts.append(elem.text or '')
    return ''.join(text_parts)


def add_comment(doc_path: str, para_index: int, comment_text: str, author: str = "LLM Agent",
                initials: str = "LA", start_char: Optional[int] = None, end_char: Optional[int] = None, output_path: Optional[str] = None) -> str:
    """Add a comment to a paragraph or specific text range."""
    doc, err = load_document(doc_path)
    if err:
        return err

    if para_index >= len(doc.paragraphs):
        return errors.err("annotations", "add_comment", "Invalid paragraph index.")

    # Get or create comments part
    comments_part = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith('/comments'):
            comments_part = rel.target_part
            break

    # Find the next available comment ID
    existing_ids = set()
    body = doc.element.body
    for cs in body.iter(qn('w:commentRangeStart')):
        existing_ids.add(int(cs.get(qn('w:id'), '0')))
    comment_id = str(max(existing_ids, default=-1) + 1)

    now = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    # If no comments part exists, create one
    if comments_part is None:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        comments_xml_bytes = (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:comment w:id="{comment_id}" w:author="{author}" w:date="{now}" w:initials="{initials}">'
            f'<w:p><w:r><w:t xml:space="preserve">{comment_text}</w:t></w:r></w:p>'
            f'</w:comment></w:comments>'
        ).encode('utf-8')
        comments_part = Part(
            PackURI('/word/comments.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
            comments_xml_bytes,
            doc.part.package
        )
        doc.part.relate_to(comments_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
    else:
        # Add to existing comments
        comment_elem = parse_xml(
            f'<w:comment {nsdecls("w")} w:id="{comment_id}" w:author="{author}" w:date="{now}" w:initials="{initials}">'
            f'<w:p><w:r><w:t xml:space="preserve">{comment_text}</w:t></w:r></w:p>'
            f'</w:comment>'
        )
        comments_part._element.append(comment_elem)

    # Add comment range markers to the paragraph
    para = doc.paragraphs[para_index]
    range_start = parse_xml(f'<w:commentRangeStart {nsdecls("w")} w:id="{comment_id}"/>')
    range_end = parse_xml(f'<w:commentRangeEnd {nsdecls("w")} w:id="{comment_id}"/>')
    comment_ref = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
        f'<w:commentReference w:id="{comment_id}"/></w:r>'
    )

    para._p.insert(0, range_start)
    para._p.append(range_end)
    para._p.append(comment_ref)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Comment #{comment_id} added (P{para_index}).")


def delete_comment(doc_path: str, comment_id: Union[int, str], output_path: Optional[str] = None) -> str:
    """Delete a comment by its ID."""
    doc, err = load_document(doc_path)
    if err:
        return err

    # Remove from comments part
    for rel in doc.part.rels.values():
        if rel.reltype.endswith('/comments'):
            comments_xml = rel.target_part._element
            for comment in comments_xml.findall(qn('w:comment')):
                if comment.get(qn('w:id')) == str(comment_id):
                    comments_xml.remove(comment)
                    break

    # Remove range markers from body
    body = doc.element.body
    for tag in ['commentRangeStart', 'commentRangeEnd']:
        for elem in body.iter(qn(f'w:{tag}')):
            if elem.get(qn('w:id')) == str(comment_id):
                elem.getparent().remove(elem)

    # Remove comment references
    for ref in body.iter(qn('w:commentReference')):
        if ref.get(qn('w:id')) == str(comment_id):
            run = ref.getparent()
            if run is not None:
                run.getparent().remove(run)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Comment #{comment_id} deleted.")


# ===================== FOOTNOTES & ENDNOTES =====================

def read_footnotes(doc_path: str) -> str:
    """Read all footnotes from the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    footnotes_part = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith('/footnotes'):
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        return "No footnotes found in the document."

    footnotes_xml = footnotes_part._element
    footnotes = footnotes_xml.findall(qn('w:footnote'))

    lines = ["=== FOOTNOTES ==="]
    for fn in footnotes:
        fn_id = fn.get(qn('w:id'), '')
        fn_type = fn.get(qn('w:type'), 'normal')
        if fn_type in ('separator', 'continuationSeparator'):
            continue
        text_parts = []
        for p in fn.findall(qn('w:p')):
            runs = p.findall(f'.//{qn("w:t")}')
            text = ''.join(r.text or '' for r in runs)
            text_parts.append(text)
        lines.append(f"  [Footnote #{fn_id}] {' '.join(text_parts)}")

    return "\n".join(lines) if len(lines) > 1 else "No footnotes found in the document."


def read_endnotes(doc_path: str) -> str:
    """Read all endnotes from the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    endnotes_part = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith('/endnotes'):
            endnotes_part = rel.target_part
            break

    if endnotes_part is None:
        return "No endnotes found in the document."

    endnotes_xml = endnotes_part._element
    endnotes = endnotes_xml.findall(qn('w:endnote'))

    lines = ["=== ENDNOTES ==="]
    for en in endnotes:
        en_id = en.get(qn('w:id'), '')
        en_type = en.get(qn('w:type'), 'normal')
        if en_type in ('separator', 'continuationSeparator'):
            continue
        text_parts = []
        for p in en.findall(qn('w:p')):
            runs = p.findall(f'.//{qn("w:t")}')
            text = ''.join(r.text or '' for r in runs)
            text_parts.append(text)
        lines.append(f"  [Endnote #{en_id}] {' '.join(text_parts)}")

    return "\n".join(lines) if len(lines) > 1 else "No endnotes found in the document."


def add_footnote(doc_path: str, para_index: int, footnote_text: str, output_path: Optional[str] = None) -> str:
    """Add a footnote to a paragraph."""
    doc, err = load_document(doc_path)
    if err:
        return err

    if para_index >= len(doc.paragraphs):
        return errors.err("annotations", "add_footnote", "Invalid paragraph index.")

    # Get the footnotes part
    footnotes_part = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith('/footnotes'):
            footnotes_part = rel.target_part
            break

    # Find next footnote ID
    existing_ids = set()
    body = doc.element.body
    for fn_ref in body.iter(qn('w:footnoteReference')):
        existing_ids.add(int(fn_ref.get(qn('w:id'), '0')))
    fn_id = str(max(existing_ids, default=0) + 1)

    if footnotes_part is not None:
        fn_elem = parse_xml(
            f'<w:footnote {nsdecls("w")} w:id="{fn_id}">'
            f'<w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
            f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
            f'<w:footnoteRef/></w:r>'
            f'<w:r><w:t xml:space="preserve"> {footnote_text}</w:t></w:r></w:p>'
            f'</w:footnote>'
        )
        footnotes_part._element.append(fn_elem)

    # Add reference in the paragraph
    para = doc.paragraphs[para_index]
    fn_ref_run = parse_xml(
        f'<w:r {nsdecls("w")}><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        f'<w:footnoteReference w:id="{fn_id}"/></w:r>'
    )
    para._p.append(fn_ref_run)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Footnote #{fn_id} added (P{para_index}).")


# ===================== TRACKED CHANGES =====================

def read_tracked_changes(doc_path: str) -> str:
    """Read all tracked changes (insertions, deletions, formatting changes)."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body = doc.element.body
    changes = []

    # Insertions
    for ins in body.iter(qn('w:ins')):
        author = ins.get(qn('w:author'), '(unknown)')
        date = ins.get(qn('w:date'), '(no date)')
        texts = []
        for t in ins.iter(qn('w:t')):
            texts.append(t.text or '')
        text = ''.join(texts)
        changes.append(f"  [INSERTED] Author: {author} | Date: {date}\n    Text: \"{text}\"")

    # Deletions
    for dl in body.iter(qn('w:del')):
        author = dl.get(qn('w:author'), '(unknown)')
        date = dl.get(qn('w:date'), '(no date)')
        texts = []
        for t in dl.iter(qn('w:delText')):
            texts.append(t.text or '')
        text = ''.join(texts)
        changes.append(f"  [DELETED] Author: {author} | Date: {date}\n    Text: \"{text}\"")

    # Property changes
    for rpc in body.iter(qn('w:rPrChange')):
        author = rpc.get(qn('w:author'), '(unknown)')
        date = rpc.get(qn('w:date'), '(no date)')
        changes.append(f"  [FORMAT CHANGE] Author: {author} | Date: {date}")

    # Paragraph property changes
    for ppc in body.iter(qn('w:pPrChange')):
        author = ppc.get(qn('w:author'), '(unknown)')
        date = ppc.get(qn('w:date'), '(no date)')
        changes.append(f"  [PARAGRAPH CHANGE] Author: {author} | Date: {date}")

    if not changes:
        return "No tracked changes found in the document."
    return f"=== TRACKED CHANGES ({len(changes)} items) ===\n" + "\n\n".join(changes)


def accept_all_changes(doc_path: str, output_path: Optional[str] = None) -> str:
    """Accept all tracked changes in the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body = doc.element.body
    count = 0

    # Accept insertions — unwrap the content
    for ins in list(body.iter(qn('w:ins'))):
        parent = ins.getparent()
        for child in list(ins):
            ins.addprevious(child)
        parent.remove(ins)
        count += 1

    # Accept deletions — remove the content
    for dl in list(body.iter(qn('w:del'))):
        dl.getparent().remove(dl)
        count += 1

    # Remove property change tracking
    for rpc in list(body.iter(qn('w:rPrChange'))):
        rpc.getparent().remove(rpc)
        count += 1

    for ppc in list(body.iter(qn('w:pPrChange'))):
        ppc.getparent().remove(ppc)
        count += 1

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"{count} changes accepted.")


def reject_all_changes(doc_path: str, output_path: Optional[str] = None) -> str:
    """Reject all tracked changes in the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body = doc.element.body
    count = 0

    # Reject insertions — remove the content
    for ins in list(body.iter(qn('w:ins'))):
        ins.getparent().remove(ins)
        count += 1

    # Reject deletions — unwrap the deleted content back
    for dl in list(body.iter(qn('w:del'))):
        parent = dl.getparent()
        for child in list(dl):
            # Convert delText back to t
            for dt in child.iter(qn('w:delText')):
                dt.tag = qn('w:t')
            dl.addprevious(child)
        parent.remove(dl)
        count += 1

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"{count} changes rejected.")


# ===================== BOOKMARKS =====================

def list_bookmarks(doc_path: str) -> str:
    """List all bookmarks in the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body = doc.element.body
    bookmarks = []
    for bm_start in body.iter(qn('w:bookmarkStart')):
        bm_id = bm_start.get(qn('w:id'), '')
        bm_name = bm_start.get(qn('w:name'), '')
        if bm_name.startswith('_'):
            continue  # Skip internal bookmarks
        bookmarks.append(f"  [BM#{bm_id}] {bm_name}")

    if not bookmarks:
        return "No bookmarks found in the document."
    return "=== BOOKMARKS ===\n" + "\n".join(bookmarks)


def add_bookmark(doc_path: str, para_index: int, bookmark_name: str, output_path: Optional[str] = None) -> str:
    """Add a bookmark to a paragraph."""
    doc, err = load_document(doc_path)
    if err:
        return err

    if para_index >= len(doc.paragraphs):
        return errors.err("annotations", "add_bookmark", "Invalid paragraph index.")

    # Find next bookmark ID
    body = doc.element.body
    existing_ids = set()
    for bm in body.iter(qn('w:bookmarkStart')):
        existing_ids.add(int(bm.get(qn('w:id'), '0')))
    bm_id = str(max(existing_ids, default=-1) + 1)

    para = doc.paragraphs[para_index]
    bm_start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="{bookmark_name}"/>')
    bm_end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')

    para._p.insert(0, bm_start)
    para._p.append(bm_end)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Bookmark added: '{bookmark_name}' (P{para_index}).")


# ===================== TEXT BOXES & SHAPES =====================

def read_textboxes(doc_path: str) -> str:
    """Read all text boxes and shapes in the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body = doc.element.body
    textboxes = []

    # Standard textboxes in shapes
    for i, txbx in enumerate(body.iter(qn('w:txbxContent'))):
        texts = []
        for p in txbx.findall(qn('w:p')):
            runs = p.findall(f'.//{qn("w:t")}')
            text = ''.join(r.text or '' for r in runs)
            texts.append(text)
        textboxes.append(f"  [TextBox#{i}] {' | '.join(texts)}")

    # Also check fallback textboxes (wps namespace)
    wps_ns = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    for i, shape in enumerate(body.iter(f'{{{wps_ns}}}txbx')):
        for txbx_content in shape.iter(qn('w:txbxContent')):
            texts = []
            for p in txbx_content.findall(qn('w:p')):
                runs = p.findall(f'.//{qn("w:t")}')
                text = ''.join(r.text or '' for r in runs)
                texts.append(text)
            if texts:
                textboxes.append(f"  [Shape#{i}] {' | '.join(texts)}")

    if not textboxes:
        return "No text boxes found in the document."
    return "=== TEXT BOXES ===\n" + "\n".join(textboxes)
