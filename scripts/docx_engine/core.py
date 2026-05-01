"""Core document operations: create, open, save, info, metadata."""
import logging
import os
import zipfile
from typing import Any, Dict, Optional, Tuple, Union

from docx import Document
from docx.document import Document as DocumentType
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn

from docx_engine import errors

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
    'dc': 'http://purl.org/dc/elements/1.1/',
    'dcterms': 'http://purl.org/dc/terms/',
}

# Configure logging
logger = logging.getLogger("docx_engine")
if not logger.handlers:
    _handler = logging.StreamHandler()
    _formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    _handler.setFormatter(_formatter)
    logger.addHandler(_handler)
    logger.setLevel(logging.WARNING)


def load_document(doc_path: str) -> Tuple[Optional[DocumentType], Optional[str]]:
    if not os.path.exists(doc_path):
        return None, errors.err("core", "load_document", f"File '{doc_path}' not found.")
    try:
        doc = Document(doc_path)
        return doc, None
    except PackageNotFoundError:
        return None, errors.err("core", "load_document", f"File '{doc_path}' is not a valid DOCX or is encrypted.")
    except zipfile.BadZipFile:
        return None, errors.err("core", "load_document", f"File '{doc_path}' is corrupted (not a valid ZIP).")
    except Exception as e:
        return None, errors.err("core", "load_document", f"Could not open file: {e}")


def create_document(doc_path: str) -> str:
    doc = Document()
    doc.save(doc_path)
    return errors.ok(f"New blank document created: {doc_path}")


def save_document(doc: DocumentType, doc_path: str, output_path: Optional[str] = None) -> str:
    save_path = output_path if output_path else doc_path
    doc.save(save_path)
    return errors.ok(f"Document saved: {save_path}")


def get_info(doc_path: str, json_mode: bool = False) -> Union[str, Dict[str, Any]]:
    doc, err = load_document(doc_path)
    if err:
        return err

    props = doc.core_properties
    body = doc.element.body

    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    char_count = sum(len(p.text) for p in doc.paragraphs)

    image_count = len(body.findall('.//wp:inline', NSMAP)) + len(body.findall('.//wp:anchor', NSMAP))
    comment_count = len(body.findall('.//w:commentRangeStart', NSMAP))
    footnote_count = len(body.findall('.//w:footnoteReference', NSMAP))
    endnote_count = len(body.findall('.//w:endnoteReference', NSMAP))
    section_count = len(doc.sections)
    has_headers = any(not s.header.is_linked_to_previous for s in doc.sections if s.header)
    has_footers = any(not s.footer.is_linked_to_previous for s in doc.sections if s.footer)

    info = {
        "file": doc_path,
        "file_size_bytes": os.path.getsize(doc_path),
        "title": props.title or "(none)",
        "author": props.author or "(none)",
        "subject": props.subject or "(none)",
        "keywords": props.keywords or "(none)",
        "created": str(props.created) if props.created else "(none)",
        "modified": str(props.modified) if props.modified else "(none)",
        "last_modified_by": props.last_modified_by or "(none)",
        "revision": props.revision or 0,
        "category": props.category or "(none)",
        "paragraph_count": para_count,
        "table_count": table_count,
        "word_count": word_count,
        "char_count": char_count,
        "image_count": image_count,
        "comment_count": comment_count,
        "footnote_count": footnote_count,
        "endnote_count": endnote_count,
        "section_count": section_count,
        "has_headers": has_headers,
        "has_footers": has_footers,
    }

    if json_mode:
        return {"status": "SUCCESS", "info": info}

    lines = ["=== DOCUMENT INFORMATION ==="]
    for k, v in info.items():
        lines.append(f"  {k}: {v}")
    return "\n".join(lines)


def set_metadata(doc_path: str, output_path: Optional[str] = None, **kwargs: Any) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    props = doc.core_properties
    field_map = {
        'title': 'title', 'author': 'author', 'subject': 'subject',
        'keywords': 'keywords', 'category': 'category',
        'comments': 'comments', 'last_modified_by': 'last_modified_by',
    }

    updated = []
    for key, val in kwargs.items():
        if key in field_map and val is not None:
            setattr(props, field_map[key], val)
            updated.append(key)

    if not updated:
        return errors.warn("core", "set_metadata", "No metadata fields specified for update.")

    return save_document(doc, doc_path, output_path) + f"\nUpdated fields: {', '.join(updated)}"


def has_page_break(paragraph: Any) -> bool:
    """Detect if a paragraph contains a hard page break."""
    if paragraph.paragraph_format.page_break_before:
        return True

    # Check runs for <w:br w:type="page"/>
    for run in paragraph.runs:
        for br in run._r.findall(f'.//{qn("w:br")}'):
            if br.get(qn('w:type')) == 'page':
                return True
    return False
