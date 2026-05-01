"""Visual rendering: convert DOCX pages to images so the LLM can 'see' the document."""
import os
import subprocess
import tempfile
import shutil
from docx_engine.core import load_document


def render_pages(doc_path, pages=None, output_dir=None, dpi=200):
    """Convert DOCX to images via LibreOffice -> PDF -> PNG pipeline.
    
    Args:
        doc_path: Path to the .docx file
        pages: Optional list of page numbers (1-indexed), or None for all
        output_dir: Where to save PNGs (defaults to same directory as doc)
        dpi: Resolution (default 200)
    
    Returns:
        List of image paths, or error message.
    """
    if not os.path.exists(doc_path):
        return f"ERROR: '{doc_path}' not found."

    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(doc_path))

    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(doc_path))[0]

    # Step 1: Convert DOCX to PDF using LibreOffice
    pdf_path = _convert_to_pdf(doc_path, output_dir)
    if pdf_path is None:
        return _fallback_text_render(doc_path, pages)

    # Step 2: Convert PDF pages to PNG images
    try:
        from pdf2image import convert_from_path
        
        kwargs = {'dpi': dpi}
        if pages:
            kwargs['first_page'] = min(pages)
            kwargs['last_page'] = max(pages)
        
        images = convert_from_path(pdf_path, **kwargs)
        
        result_paths = []
        for i, img in enumerate(images):
            page_num = (min(pages) + i) if pages else (i + 1)
            if pages and page_num not in pages:
                continue
            img_path = os.path.join(output_dir, f"{base_name}_page_{page_num}.png")
            img.save(img_path, 'PNG')
            result_paths.append(img_path)
        
        # Clean up PDF
        try:
            os.remove(pdf_path)
        except Exception:
            pass
        
        lines = [f"=== RENDER RESULTS ({len(result_paths)} pages) ==="]
        for p in result_paths:
            lines.append(f"  PAGE: {p}")
        lines.append("\nExamine these image files using your image viewing tool.")
        return "\n".join(lines)
        
    except ImportError:
        return ("ERROR: pdf2image library not found.\n"
                "Installation: pip install pdf2image\n"
                "Poppler is also required: https://github.com/oschwartz10612/poppler-windows/releases\n\n"
                + _fallback_text_render(doc_path, pages))
    except Exception as e:
        return f"ERROR: PDF to image conversion failed: {e}\n\n" + _fallback_text_render(doc_path, pages)


def _convert_to_pdf(doc_path, output_dir):
    """Convert DOCX to PDF using LibreOffice."""
    import shutil as _shutil
    # Try common LibreOffice paths on Windows + PATH lookup
    lo_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    lo_exe = _shutil.which("soffice") or _shutil.which("soffice.exe")
    if lo_exe is None:
        for path in lo_paths:
            if os.path.exists(path):
                lo_exe = path
                break

    if lo_exe is None:
        return None

    try:
        cmd = [
            lo_exe, '--headless', '--convert-to', 'pdf',
            '--outdir', output_dir, doc_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        return None
    except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
        return None


def _fallback_text_render(doc_path, pages=None):
    """Create a detailed text representation of the document layout when visual rendering is unavailable."""
    doc, err = load_document(doc_path)
    if err:
        return err

    lines = ["=== TEXT-BASED LAYOUT VIEW ==="]
    lines.append("(Since LibreOffice was not found, text layout is shown instead of visual rendering)\n")

    # Section info
    for s_idx, section in enumerate(doc.sections):
        w = section.page_width
        h = section.page_height
        orient = "Landscape" if section.orientation else "Portrait"
        lines.append(f"--- Section {s_idx}: {orient} | Size: {w}x{h} emu ---")
        lines.append(f"    Margins: top={section.top_margin}, bottom={section.bottom_margin}, "
                     f"left={section.left_margin}, right={section.right_margin}")
        
        # Header
        if section.header and not section.header.is_linked_to_previous:
            h_text = " | ".join(p.text for p in section.header.paragraphs if p.text.strip())
            if h_text:
                lines.append(f"    [HEADER] {h_text}")
        
        # Footer
        if section.footer and not section.footer.is_linked_to_previous:
            f_text = " | ".join(p.text for p in section.footer.paragraphs if p.text.strip())
            if f_text:
                lines.append(f"    [FOOTER] {f_text}")

    lines.append("")

    # Approximate page simulation (~45 lines per page)
    current_page = 1
    line_count_on_page = 0
    LINES_PER_PAGE = 45

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        # Check for page breaks
        has_page_break = False
        pf = para.paragraph_format
        if pf.page_break_before:
            has_page_break = True
        for run in para.runs:
            for br in run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                    has_page_break = True

        if has_page_break and line_count_on_page > 0:
            current_page += 1
            line_count_on_page = 0
            lines.append(f"\n{'='*60}")
            lines.append(f"  === PAGE {current_page} ===")
            lines.append(f"{'='*60}")

        if pages and current_page not in pages:
            # Estimate lines for this paragraph
            estimated_lines = max(1, len(text) // 80 + 1)
            line_count_on_page += estimated_lines
            if line_count_on_page >= LINES_PER_PAGE:
                current_page += 1
                line_count_on_page = 0
            continue

        if line_count_on_page == 0:
            lines.append(f"\n{'='*60}")
            lines.append(f"  === PAGE {current_page} ===")
            lines.append(f"{'='*60}")

        if style.startswith('Heading'):
            level = style.replace('Heading ', '').strip()
            lines.append(f"  [P{i}] {'#' * int(level) if level.isdigit() else '#'} {text}")
        elif text:
            lines.append(f"  [P{i}] {text}")
        else:
            lines.append(f"  [P{i}] <empty line>")

        estimated_lines = max(1, len(text) // 80 + 1)
        line_count_on_page += estimated_lines
        if line_count_on_page >= LINES_PER_PAGE:
            current_page += 1
            line_count_on_page = 0

    lines.append(f"\n--- Total estimated page count: {current_page} ---")
    return "\n".join(lines)


def describe_layout(doc_path):
    """Provide a comprehensive layout description for the LLM to understand the visual structure."""
    doc, err = load_document(doc_path)
    if err:
        return err

    from docx.oxml.ns import qn as _qn

    lines = ["=== LAYOUT ANALYSIS ==="]

    for s_idx, section in enumerate(doc.sections):
        lines.append(f"\n--- Section {s_idx} ---")
        
        # Page dimensions in cm
        def emu_to_cm(emu):
            return round(emu / 360000, 2) if emu else 0

        lines.append(f"  Page: {emu_to_cm(section.page_width)}x{emu_to_cm(section.page_height)} cm")
        lines.append(f"  Orientation: {'Landscape' if section.orientation else 'Portrait'}")
        lines.append(f"  Margins (cm): top={emu_to_cm(section.top_margin)}, "
                     f"bottom={emu_to_cm(section.bottom_margin)}, "
                     f"left={emu_to_cm(section.left_margin)}, right={emu_to_cm(section.right_margin)}")

        # Header/footer info
        if section.header:
            linked = section.header.is_linked_to_previous
            h_texts = [p.text for p in section.header.paragraphs if p.text.strip()]
            lines.append(f"  Header: {'(linked)' if linked else '|'.join(h_texts) if h_texts else '(empty)'}")

        if section.footer:
            linked = section.footer.is_linked_to_previous
            f_texts = [p.text for p in section.footer.paragraphs if p.text.strip()]
            lines.append(f"  Footer: {'(linked)' if linked else '|'.join(f_texts) if f_texts else '(empty)'}")

        # Different first page header/footer
        if section.different_first_page_header_footer:
            lines.append(f"  Different first page header/footer: Yes")

    # Content structure overview
    lines.append(f"\n--- Content Structure ---")
    
    # Count element types
    body = doc.element.body
    inline_imgs = len(list(body.iter(_qn('wp:inline'))))
    anchor_imgs = len(list(body.iter(_qn('wp:anchor'))))
    textboxes = len(list(body.iter(_qn('w:txbxContent'))))
    
    lines.append(f"  Paragraph count: {len(doc.paragraphs)}")
    lines.append(f"  Table count: {len(doc.tables)}")
    lines.append(f"  Inline image: {inline_imgs}")
    lines.append(f"  Anchored image: {anchor_imgs}")
    lines.append(f"  Text box: {textboxes}")

    return "\n".join(lines)
