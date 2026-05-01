"""Extended elements: SmartArt, Charts, OLE objects, content controls, protection, watermarks."""
import os
import base64
from typing import List, Optional, Dict, Any, Union, Tuple
from lxml import etree
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Cm, RGBColor
from docx_engine.core import load_document, save_document
from docx_engine import errors


# Namespace URIs not in python-docx's built-in map
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'dgm': 'http://schemas.openxmlformats.org/drawingml/2006/diagram',
    'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'o': 'urn:schemas-microsoft-com:office:office',
    'v': 'urn:schemas-microsoft-com:vml',
    'w10': 'urn:schemas-microsoft-com:office:word',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'sdt': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}


# ===================== SMARTART =====================

def read_smartart(doc_path: str) -> str:
    """Detect and describe SmartArt diagrams in the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    smartarts = []
    
    # Check relationships for diagram parts
    for i, rel in enumerate(doc.part.rels.values()):
        if 'diagramData' in rel.reltype or 'diagramLayout' in rel.reltype:
            smartarts.append(f"  [SmartArt rel] {rel.reltype.split('/')[-1]}: {rel.target_ref}")

    # Check for dgm:relIds in document body
    body = doc.element.body
    dgm_ns = 'http://schemas.openxmlformats.org/drawingml/2006/diagram'
    for i, elem in enumerate(body.iter(f'{{{dgm_ns}}}relIds')):
        attrs = {k.split('}')[-1]: v for k, v in elem.attrib.items()}
        smartarts.append(f"  [SmartArt#{i}] DiagramData: {attrs}")

    # Also check for graphicData with diagram URI
    for i, gd in enumerate(body.iter(f'{{{NS["a"]}}}graphicData')):
        uri = gd.get('uri', '')
        if 'diagram' in uri.lower():
            # Try to extract text from the diagram
            texts = []
            for t in gd.iter(f'{{{NS["a"]}}}t'):
                if t.text:
                    texts.append(t.text)
            text_preview = " | ".join(texts[:10]) if texts else "(text could not be extracted)"
            smartarts.append(f"  [Diagram#{i}] URI: {uri}\n    Text: {text_preview}")

    if not smartarts:
        return errors.err("extended", "read_smartart", "No SmartArt/Diagram found in document.")
    return "=== SMARTART / DIAGRAMS ===\n" + "\n".join(smartarts)


# ===================== CHARTS =====================

def read_charts(doc_path: str) -> str:
    """Detect and describe embedded charts."""
    doc, err = load_document(doc_path)
    if err:
        return err

    charts = []
    
    # Check relationships for chart parts
    for i, rel in enumerate(doc.part.rels.values()):
        if 'chart' in rel.reltype.lower():
            chart_info = f"  [Chart rel] {rel.target_ref}"
            
            # Try to read chart data from the part
            try:
                chart_part = rel.target_part
                chart_xml = chart_part._element
                
                # Extract chart title
                chart_ns = 'http://schemas.openxmlformats.org/drawingml/2006/chart'
                titles = []
                for t in chart_xml.iter(f'{{{NS["a"]}}}t'):
                    if t.text:
                        titles.append(t.text)
                
                if titles:
                    chart_info += f"\n    Title/Labels: {' | '.join(titles[:20])}"
                
                # Try to identify chart type
                chart_types = []
                for tag in ['barChart', 'lineChart', 'pieChart', 'areaChart', 'scatterChart',
                           'doughnutChart', 'radarChart', 'bubbleChart', 'surfaceChart']:
                    if chart_xml.find(f'.//{{{chart_ns}}}{tag}') is not None:
                        chart_types.append(tag)
                
                if chart_types:
                    chart_info += f"\n    Chart type: {', '.join(chart_types)}"
                    
            except Exception as e:
                chart_info += f"\n    (Details could not be read: {e})"
            
            charts.append(chart_info)

    # Check for chart references in body
    body = doc.element.body
    chart_ns = 'http://schemas.openxmlformats.org/drawingml/2006/chart'
    for gd in body.iter(f'{{{NS["a"]}}}graphicData'):
        uri = gd.get('uri', '')
        if 'chart' in uri.lower():
            charts.append(f"  [Chart inline] URI: {uri}")

    if not charts:
        return errors.err("extended", "read_charts", "No charts found in document.")
    return "=== CHARTS ===\n" + "\n".join(charts)


# ===================== EMBEDDED / OLE OBJECTS =====================

def list_embedded_objects(doc_path: str) -> str:
    """List all embedded OLE objects (Excel, PDF, etc.)."""
    doc, err = load_document(doc_path)
    if err:
        return err

    objects = []
    
    # Check relationships for OLE objects
    for i, rel in enumerate(doc.part.rels.values()):
        reltype = rel.reltype.lower()
        if 'oleobject' in reltype or 'package' in reltype or 'embeddings' in reltype:
            objects.append(f"  [OLE#{i}] Type: {rel.reltype.split('/')[-1]} | Target: {rel.target_ref}")

    # Check for OLE objects in body XML
    body = doc.element.body
    ole_ns = 'urn:schemas-microsoft-com:office:office'
    for i, ole in enumerate(body.iter(f'{{{ole_ns}}}OLEObject')):
        prog_id = ole.get('ProgID', '(unknown)')
        r_id = ole.get(f'{{{NS["r"]}}}id', '')
        objects.append(f"  [OLEObject#{i}] Program: {prog_id} | rId: {r_id}")

    # Check for embedded packages
    for i, obj in enumerate(body.iter(qn('w:object'))):
        ole_elems = obj.findall(f'{{{ole_ns}}}OLEObject')
        for ole in ole_elems:
            prog_id = ole.get('ProgID', '(unknown)')
            objects.append(f"  [EmbObj#{i}] {prog_id}")

    if not objects:
        return errors.err("extended", "list_embedded_objects", "No embedded objects found in document.")
    return "=== EMBEDDED OBJECTS (OLE) ===\n" + "\n".join(objects)


def extract_embedded(doc_path: str, rel_index: int, output_path: str) -> str:
    """Extract an embedded object by its relationship index."""
    doc, err = load_document(doc_path)
    if err:
        return err

    rels = list(doc.part.rels.values())
    matching = []
    for i, rel in enumerate(rels):
        reltype = rel.reltype.lower()
        if 'oleobject' in reltype or 'package' in reltype or 'embeddings' in reltype:
            matching.append((i, rel))

    if rel_index >= len(matching):
        return errors.err("extended", "extract_embedded", f"Invalid embedded object index. {len(matching)} objects exist.")

    _, rel = matching[rel_index]
    try:
        part = rel.target_part
        with open(output_path, 'wb') as f:
            f.write(part.blob)
        return errors.ok(f"Embedded object extracted → {output_path} ({len(part.blob)} bytes)")
    except Exception as e:
        return errors.err("extended", "extract_embedded", f"Object could not be extracted: {e}")


# ===================== CONTENT CONTROLS (SDT) =====================

def read_content_controls(doc_path: str) -> str:
    """Read all Structured Document Tags (content controls / form fields)."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body = doc.element.body
    controls = []

    for i, sdt in enumerate(body.iter(qn('w:sdt'))):
        # Get properties
        sdt_pr = sdt.find(qn('w:sdtPr'))
        tag_elem = sdt_pr.find(qn('w:tag')) if sdt_pr is not None else None
        alias_elem = sdt_pr.find(qn('w:alias')) if sdt_pr is not None else None
        
        tag = tag_elem.get(qn('w:val'), '') if tag_elem is not None else ''
        alias = alias_elem.get(qn('w:val'), '') if alias_elem is not None else ''

        # Determine type
        sdt_type = "Unknown"
        if sdt_pr is not None:
            if sdt_pr.find(qn('w:text')) is not None:
                sdt_type = "Plain Text"
            elif sdt_pr.find(qn('w14:checkbox')) is not None:
                sdt_type = "Checkbox"
            elif sdt_pr.find(qn('w:richText')) is not None:
                sdt_type = "Rich Text"
            elif sdt_pr.find(qn('w:comboBox')) is not None:
                sdt_type = "Combo Box"
            elif sdt_pr.find(qn('w:dropDownList')) is not None:
                sdt_type = "Dropdown List"
            elif sdt_pr.find(qn('w:date')) is not None:
                sdt_type = "Date Picker"
            elif sdt_pr.find(qn('w:picture')) is not None:
                sdt_type = "Picture"
            elif sdt_pr.find(qn('w:docPartList')) is not None:
                sdt_type = "Document Part"

        # Get content
        sdt_content = sdt.find(qn('w:sdtContent'))
        content_text = ""
        if sdt_content is not None:
            texts = []
            for t in sdt_content.iter(qn('w:t')):
                if t.text:
                    texts.append(t.text)
            content_text = ''.join(texts)

        controls.append(
            f"  [SDT#{i}] Type: {sdt_type}\n"
            f"    Tag: {tag or '(none)'} | Alias: {alias or '(none)'}\n"
            f"    Content: \"{content_text[:80]}{'...' if len(content_text) > 80 else ''}\""
        )

    if not controls:
        return errors.err("extended", "read_content_controls", "No content controls (form fields) found in document.")
    return f"=== CONTENT CONTROLS ({len(controls)} items) ===\n" + "\n\n".join(controls)


# ===================== DOCUMENT PROTECTION =====================

def read_protection(doc_path: str) -> str:
    """Read document protection settings."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body_parent = doc.element
    settings = {}

    # Check document protection element
    for dp in body_parent.iter(qn('w:documentProtection')):
        settings['edit'] = dp.get(qn('w:edit'), '(none)')
        settings['enforcement'] = dp.get(qn('w:enforcement'), '0')
        settings['formatting'] = dp.get(qn('w:formatting'), '(none)')

    # Check settings part for more protection info
    for rel in doc.part.rels.values():
        if rel.reltype.endswith('/settings'):
            try:
                settings_xml = rel.target_part._element
                for dp in settings_xml.iter(qn('w:documentProtection')):
                    settings['edit'] = dp.get(qn('w:edit'), settings.get('edit', '(none)'))
                    settings['enforcement'] = dp.get(qn('w:enforcement'), settings.get('enforcement', '0'))
                    settings['cryptAlgorithm'] = dp.get(qn('w:cryptAlgorithmSid'), '(none)')
            except Exception:
                pass

    if not settings:
        return errors.ok("Document protection not applied.")

    protection_types = {
        'readOnly': 'Read Only',
        'comments': 'Comments Only',
        'trackedChanges': 'Tracked Changes Only',
        'forms': 'Form Filling Only',
        'none': 'No Protection',
    }

    edit_type = settings.get('edit', 'none')
    enforced = settings.get('enforcement', '0') == '1'

    lines = ["=== DOCUMENT PROTECTION ==="]
    lines.append(f"  Protection type: {protection_types.get(edit_type, edit_type)}")
    lines.append(f"  Is enforced: {'Yes' if enforced else 'No'}")
    if 'cryptAlgorithm' in settings:
        lines.append(f"  Encryption: {settings['cryptAlgorithm']}")

    return "\n".join(lines)


# ===================== WATERMARK =====================

def add_watermark(doc_path: str, text: str, font_size: int = 72, color: str = "#C0C0C0", rotation: int = -45, output_path: Optional[str] = None) -> str:
    """Add a diagonal text watermark to all pages."""
    doc, err = load_document(doc_path)
    if err:
        return err

    color_hex = color.lstrip('#')

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Create watermark using VML shape in header
        watermark_xml = (
            f'<w:p {nsdecls("w", "r")} xmlns:v="urn:schemas-microsoft-com:vml" '
            f'xmlns:o="urn:schemas-microsoft-com:office:office" '
            f'xmlns:w10="urn:schemas-microsoft-com:office:word">'
            f'<w:r>'
            f'<w:rPr><w:noProof/></w:rPr>'
            f'<w:pict>'
            f'<v:shapetype id="_x0000_t136" coordsize="21600,21600" '
            f'o:spt="136" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">'
            f'</v:shapetype>'
            f'<v:shape id="PowerPlusWaterMarkObject" '
            f'o:spid="_x0000_s2049" type="#_x0000_t136" '
            f'style="position:absolute;margin-left:0;margin-top:0;width:527.85pt;height:131.95pt;'
            f'rotation:{rotation};z-index:-251656192;mso-position-horizontal:center;'
            f'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
            f'mso-position-vertical-relative:margin" '
            f'o:allowincell="f" fillcolor="#{color_hex}" stroked="f">'
            f'<v:fill opacity=".5"/>'
            f'<v:textpath style="font-family:&quot;Calibri&quot;;font-size:{font_size}pt" '
            f'string="{text}"/>'
            f'<w10:wrap anchorx="margin" anchory="margin"/>'
            f'</v:shape>'
            f'</v:pict>'
            f'</w:r>'
            f'</w:p>'
        )

        watermark_elem = parse_xml(watermark_xml)
        header._element.append(watermark_elem)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Watermark added: '{text}'")


def remove_watermark(doc_path: str, output_path: Optional[str] = None) -> str:
    """Remove watermark from the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    removed = 0
    vml_ns = 'urn:schemas-microsoft-com:vml'
    
    for section in doc.sections:
        if section.header:
            for para in list(section.header._element.iter(qn('w:p'))):
                # Check for pict elements containing watermark shapes
                for pict in list(para.iter(qn('w:pict'))):
                    for shape in pict.iter(f'{{{vml_ns}}}shape'):
                        shape_id = shape.get('id', '')
                        if 'WaterMark' in shape_id or 'watermark' in shape_id.lower():
                            para.getparent().remove(para)
                            removed += 1
                            break

    if removed == 0:
        return errors.err("extended", "remove_watermark", "No watermark found in document.")
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"{removed} watermark(s) removed.")


# ===================== LINE NUMBERING =====================

def set_line_numbering(doc_path: str, start: int = 1, count_by: int = 1, restart: str = 'newPage', output_path: Optional[str] = None) -> str:
    """Enable line numbering for the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    restart_map = {
        'newPage': 'newPage',
        'newSection': 'newSection',
        'continuous': 'continuous',
    }

    for section in doc.sections:
        sectPr = section._sectPr
        # Remove existing line numbering
        for ln in sectPr.findall(qn('w:lnNumType')):
            sectPr.remove(ln)

        ln_num = parse_xml(
            f'<w:lnNumType {nsdecls("w")} w:start="{start}" '
            f'w:countBy="{count_by}" w:restart="{restart_map.get(restart, "newPage")}"/>'
        )
        sectPr.append(ln_num)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok("Line numbering enabled.")


# ===================== DOCUMENT STATISTICS =====================

def full_statistics(doc_path: str) -> str:
    """Comprehensive document statistics beyond basic word count."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body = doc.element.body

    # Count everything
    stats = {
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
        'sections': len(doc.sections),
        'words': sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip()),
        'characters': sum(len(p.text) for p in doc.paragraphs),
        'characters_no_spaces': sum(len(p.text.replace(' ', '')) for p in doc.paragraphs),
        'lines': sum(1 for p in doc.paragraphs if p.text.strip()),
        'empty_lines': sum(1 for p in doc.paragraphs if not p.text.strip()),
        'headings': sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading')),
        'images_inline': len(list(body.iter(f'{{{NS["wp"]}}}inline'))),
        'images_anchor': len(list(body.iter(f'{{{NS["wp"]}}}anchor'))),
        'text_boxes': len(list(body.iter(qn('w:txbxContent')))),
        'comments': len(list(body.iter(qn('w:commentRangeStart')))),
        'footnotes': len(list(body.iter(qn('w:footnoteReference')))),
        'endnotes': len(list(body.iter(qn('w:endnoteReference')))),
        'bookmarks': len([bm for bm in body.iter(qn('w:bookmarkStart'))
                         if not bm.get(qn('w:name'), '').startswith('_')]),
        'hyperlinks': len(list(body.iter(qn('w:hyperlink')))),
    }

    # Table cell count
    total_cells = 0
    for table in doc.tables:
        for row in table.rows:
            total_cells += len(row.cells)
    stats['table_cells'] = total_cells

    # Style usage
    style_usage = {}
    for p in doc.paragraphs:
        sname = p.style.name
        style_usage[sname] = style_usage.get(sname, 0) + 1

    # Font usage
    font_usage = {}
    for p in doc.paragraphs:
        for run in p.runs:
            fname = run.font.name or "(default)"
            font_usage[fname] = font_usage.get(fname, 0) + 1

    lines = ["=== DETAILED STATISTICS ==="]
    for k, v in stats.items():
        lines.append(f"  {k}: {v}")

    lines.append("\n  Style usage:")
    for style, count in sorted(style_usage.items(), key=lambda x: -x[1])[:15]:
        lines.append(f"    {style}: {count}")

    lines.append("\n  Font usage:")
    for font, count in sorted(font_usage.items(), key=lambda x: -x[1])[:10]:
        lines.append(f"    {font}: {count}")

    return "\n".join(lines)
