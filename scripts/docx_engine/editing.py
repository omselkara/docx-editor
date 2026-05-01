"""Text editing: insert, delete, replace, append paragraphs and headings."""
from __future__ import annotations

import re
from typing import Any
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx_engine.core import load_document, save_document
from docx_engine import errors


def insert_paragraph(
    doc_path: str,
    text: str,
    index: int | None = None,
    after_heading: str | None = None,
    style: str | None = None,
    output_path: str | None = None,
) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    if after_heading:
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith("Heading") and after_heading.lower() in para.text.lower():
                new_p = doc.add_paragraph(text, style=style)
                para._p.addnext(new_p._p)
                return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Text inserted after heading '{after_heading}'.")
        return errors.err("editing", "insert_paragraph", f"Heading '{after_heading}' not found.")

    if index is not None:
        paras = doc.paragraphs
        if index < 0 or index >= len(paras):
            reason = f"Invalid index {index}. Document contains {len(paras)} paragraphs (0–{len(paras)-1})."
            return errors.err("editing", "insert_paragraph", reason)
        new_p = doc.add_paragraph(text, style=style)
        paras[index]._p.addnext(new_p._p)
        return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Text inserted after P{index}.")

    doc.add_paragraph(text, style=style)
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok("Text appended to the end of the document.")


def insert_heading(
    doc_path: str,
    text: str,
    level: int = 1,
    index: int | None = None,
    output_path: str | None = None,
) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    if index is not None:
        paras = doc.paragraphs
        if index < 0 or index >= len(paras):
            reason = f"Invalid index {index}. Document contains {len(paras)} paragraphs."
            return errors.err("editing", "insert_heading", reason)
        new_p = doc.add_heading(text, level=level)
        paras[index]._p.addnext(new_p._p)
    else:
        doc.add_heading(text, level=level)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"H{level} heading added: '{text}'")


def delete_paragraphs(
    doc_path: str,
    indices: list[int] | str,
    output_path: str | None = None,
) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    paras = doc.paragraphs
    if isinstance(indices, str):
        indices = [int(x.strip()) for x in indices.split(",")]

    sorted_indices = sorted(indices, reverse=True)
    deleted = []
    for idx in sorted_indices:
        if 0 <= idx < len(paras):
            p = paras[idx]._p
            p.getparent().remove(p)
            deleted.append(idx)

    if not deleted:
        return errors.err("editing", "delete_paragraphs", "No paragraphs were deleted — all indices out of range.")
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Deleted paragraphs: {deleted}")


def replace_text(
    doc_path: str,
    find_pattern: str,
    replace_with: str,
    use_regex: bool = False,
    output_path: str | None = None,
) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    count = 0
    for para in doc.paragraphs:
        count += _surgical_replace(para, find_pattern, replace_with, use_regex)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += _surgical_replace(para, find_pattern, replace_with, use_regex)

    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for para in hf.paragraphs:
                    count += _surgical_replace(para, find_pattern, replace_with, use_regex)

    if count == 0:
        return errors.warn("editing", "replace_text", f"'{find_pattern}' not found, no changes made.")
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Replaced in {count} occurrences.")


# ─── Surgical Replace Helpers ────────────────────────────────────────────────

def _build_char_map(paragraph: Any) -> tuple[str, list[tuple[int, int]]]:
    """Return (full_text, char_map) where char_map[i] = (run_idx, char_idx_in_run)."""
    full_text = ""
    char_map: list[tuple[int, int]] = []
    for r_idx, run in enumerate(paragraph.runs):
        for c_idx, _ in enumerate(run.text):
            full_text += run.text[c_idx]
            char_map.append((r_idx, c_idx))
    return full_text, char_map


def _find_matches(full_text: str, find_str: str, use_regex: bool) -> list[tuple[int, int]]:
    """Return list of (start, end) match positions in full_text."""
    matches: list[tuple[int, int]] = []
    if use_regex:
        for m in re.finditer(find_str, full_text):
            matches.append((m.start(), m.end()))
    else:
        start = 0
        while True:
            idx = full_text.find(find_str, start)
            if idx == -1:
                break
            matches.append((idx, idx + len(find_str)))
            start = idx + len(find_str)
    return matches


def _apply_single_run_replace(run: Any, match_text: str, replace_str: str) -> None:
    """Replace match_text with replace_str within a single run (first occurrence)."""
    run.text = run.text.replace(match_text, replace_str, 1)


def _apply_multi_run_replace(
    paragraph: Any,
    char_map: list[tuple[int, int]],
    start: int,
    end: int,
    replace_str: str,
) -> None:
    """Replace a match that spans multiple runs, preserving per-run formatting."""
    start_run_idx, char_idx_in_first = char_map[start]
    end_run_idx, char_idx_in_last = char_map[end - 1]

    # First run: keep prefix, append replacement
    first_run = paragraph.runs[start_run_idx]
    first_run.text = first_run.text[:char_idx_in_first] + replace_str

    # Middle runs: clear completely
    for r_idx in range(start_run_idx + 1, end_run_idx):
        paragraph.runs[r_idx].text = ""

    # Last run: keep suffix after the matched character
    last_run = paragraph.runs[end_run_idx]
    last_run.text = last_run.text[char_idx_in_last + 1:]


def _surgical_replace(paragraph: Any, find_str: str, replace_str: str, use_regex: bool = False) -> int:
    """
    Replace all occurrences of find_str in paragraph while preserving run formatting.
    Handles matches that span multiple runs. Processes from end to start so earlier
    char_map indices stay valid after each replacement.
    """
    full_text, char_map = _build_char_map(paragraph)
    if not full_text:
        return 0

    matches = _find_matches(full_text, find_str, use_regex)
    if not matches:
        return 0

    for start, end in reversed(matches):
        start_run_idx, _ = char_map[start]
        end_run_idx, _ = char_map[end - 1]

        if start_run_idx == end_run_idx:
            _apply_single_run_replace(
                paragraph.runs[start_run_idx],
                full_text[start:end],
                replace_str,
            )
        else:
            _apply_multi_run_replace(paragraph, char_map, start, end, replace_str)

    return len(matches)


def append_text(
    doc_path: str,
    text: str,
    style: str | None = None,
    output_path: str | None = None,
) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    doc.add_paragraph(text, style=style)
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok("Text appended to the end of the document.")
