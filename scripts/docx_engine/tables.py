"""Table operations: create, read, modify, format, merge, add/delete rows/columns."""
import json
from typing import List, Optional, Dict, Any, Union
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_engine.core import load_document, save_document
from docx_engine import errors


def list_tables(doc_path: str, json_mode: bool = False) -> Union[str, Dict[str, Any]]:
    doc, err = load_document(doc_path)
    if err:
        return err
    if not doc.tables:
        if json_mode:
            return {"status": "WARNING", "tables": [], "message": "No tables found in the document."}
        return errors.warn("tables", "list_tables", "No tables found in the document.")

    tables_data = []
    lines = ["=== TABLES ==="]
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        first_row = " | ".join(cell.text[:30] for cell in table.rows[0].cells) if rows > 0 else ""
        tables_data.append({"index": i, "rows": rows, "cols": cols, "preview": first_row})
        lines.append(f"  [T{i}] {rows}x{cols} — First row: {first_row}")

    if json_mode:
        return {"status": "SUCCESS", "tables": tables_data}
    return "\n".join(lines)


def read_table(doc_path: str, table_index: int, json_mode: bool = False) -> Union[str, Dict[str, Any]]:
    doc, err = load_document(doc_path)
    if err:
        return err
    if table_index < 0 or table_index >= len(doc.tables):
        reason = f"Invalid table index. {len(doc.tables)} tables available."
        if json_mode:
            return {"status": "ERROR", "message": errors.err("tables", "read_table", reason)}
        return errors.err("tables", "read_table", reason)

    table = doc.tables[table_index]
    rows_data = []
    for row in table.rows:
        rows_data.append([cell.text.replace('\n', ' ').strip() for cell in row.cells])

    if json_mode:
        return {
            "status": "SUCCESS",
            "index": table_index,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "data": rows_data,
        }

    lines = [f"=== TABLE T{table_index} ({len(table.rows)}x{len(table.columns)}) ==="]
    # Markdown table
    if rows_data:
        headers = rows_data[0]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows_data[1:]:
            lines.append("| " + " | ".join(row) + " |")

    # Cell coordinates
    lines.append("\n--- Cell Coordinates ---")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()[:50]
            lines.append(f"  [{r_idx},{c_idx}] {text}")
    return "\n".join(lines)


def create_table(doc_path: str, headers: str, rows_data: Optional[str] = None, style: Optional[str] = None, after_para: Optional[int] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    header_list = [h.strip() for h in headers.split(",")]
    cols = len(header_list)

    data_rows = []
    if rows_data:
        for row_str in rows_data.split(";"):
            data_rows.append([c.strip() for c in row_str.split(",")])

    total_rows = 1 + len(data_rows)
    table = doc.add_table(rows=total_rows, cols=cols)

    if style:
        try:
            table.style = style
        except Exception:
            table.style = 'Table Grid'
    else:
        table.style = 'Table Grid'

    for i, header in enumerate(header_list):
        table.rows[0].cells[i].text = header

    for r_idx, row_data in enumerate(data_rows):
        for c_idx, val in enumerate(row_data):
            if c_idx < cols:
                table.rows[r_idx + 1].cells[c_idx].text = val

    if after_para is not None and 0 <= after_para < len(doc.paragraphs):
        doc.paragraphs[after_para]._p.addnext(table._tbl)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"{total_rows}x{cols} table created.")


def modify_cell(doc_path: str, table_index: int, row: int, col: int, text: str, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if table_index >= len(doc.tables):
        return errors.err("tables", "modify_cell", "Invalid table index.")

    table = doc.tables[table_index]
    if row >= len(table.rows) or col >= len(table.columns):
        return errors.err("tables", "modify_cell", f"Invalid cell coordinates [{row},{col}].")

    table.rows[row].cells[col].text = text
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"[T{table_index}][{row},{col}] updated.")


def add_row(doc_path: str, table_index: int, values: Optional[str] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if table_index >= len(doc.tables):
        return errors.err("tables", "add_row", "Invalid table index.")
    table = doc.tables[table_index]
    row = table.add_row()
    if values:
        val_list = [v.strip() for v in values.split(",")]
        for i, val in enumerate(val_list):
            if i < len(row.cells):
                row.cells[i].text = val
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Row added (T{table_index}).")


def add_column(doc_path: str, table_index: int, header: str = "", output_path: Optional[str] = None) -> str:
    """Add a column to a table by cloning the last cell in each row via lxml."""
    doc, err = load_document(doc_path)
    if err:
        return err
    if table_index >= len(doc.tables):
        return errors.err("tables", "add_column", "Invalid table index.")
    table = doc.tables[table_index]

    for r_idx, row in enumerate(table.rows):
        # Clone the last cell's XML element as template, clear its content
        last_tc = row.cells[-1]._tc
        new_tc = etree.fromstring(etree.tostring(last_tc))  # deep copy
        # Clear all paragraph text in the clone
        for t_elem in new_tc.iter(qn('w:t')):
            t_elem.text = ''
        row._tr.append(new_tc)

    # Set header text in first row's new cell
    if header:
        table.rows[0].cells[-1].text = header

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Column added (T{table_index}).")


def delete_row(doc_path: str, table_index: int, row_index: int, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if table_index >= len(doc.tables):
        return errors.err("tables", "delete_row", "Invalid table index.")
    table = doc.tables[table_index]
    if row_index >= len(table.rows):
        return errors.err("tables", "delete_row", "Invalid row index.")
    tr = table.rows[row_index]._tr
    table._tbl.remove(tr)
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Row {row_index} deleted (T{table_index}).")


def delete_column(doc_path: str, table_index: int, col_index: int, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if table_index >= len(doc.tables):
        return errors.err("tables", "delete_column", "Invalid table index.")
    table = doc.tables[table_index]
    if col_index >= len(table.columns):
        return errors.err("tables", "delete_column", "Invalid column index.")
    for row in table.rows:
        cell = row.cells[col_index]
        row._tr.remove(cell._tc)
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Column {col_index} deleted (T{table_index}).")


def merge_cells(doc_path: str, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if table_index >= len(doc.tables):
        return errors.err("tables", "merge_cells", "Invalid table index.")
    table = doc.tables[table_index]
    start_cell = table.cell(start_row, start_col)
    end_cell = table.cell(end_row, end_col)
    start_cell.merge(end_cell)
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Cells merged [{start_row},{start_col}]-[{end_row},{end_col}].")


def format_table_cell(doc_path: str, table_index: int, row: int, col: int, bg_color: Optional[str] = None,
                      bold: Optional[bool] = None, alignment: Optional[str] = None, font_size: Optional[float] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if table_index >= len(doc.tables):
        return errors.err("tables", "format_table_cell", "Invalid table index.")
    table = doc.tables[table_index]
    if row >= len(table.rows) or col >= len(table.columns):
        return errors.err("tables", "format_table_cell", f"Invalid cell coordinates [{row},{col}].")
    cell = table.rows[row].cells[col]

    if bg_color:
        color = bg_color.lstrip('#')
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        tc_pr = cell._tc.get_or_add_tcPr()
        # Remove existing shading
        for existing in tc_pr.findall(qn('w:shd')):
            tc_pr.remove(existing)
        tc_pr.append(shading)

    for para in cell.paragraphs:
        if alignment:
            align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                         'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
            if alignment.lower() in align_map:
                para.paragraph_format.alignment = align_map[alignment.lower()]
        for run in para.runs:
            if bold is not None:
                run.bold = bold
            if font_size:
                run.font.size = Pt(font_size)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Cell [{row},{col}] formatted (T{table_index}).")
