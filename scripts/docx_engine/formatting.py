"""Character and paragraph formatting."""
import re
import functools
from typing import List, Optional, Dict, Any, Union
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx_engine.core import load_document, save_document
from docx_engine import errors

ALIGN_MAP = {
    'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def format_text(doc_path: str, match: Optional[str] = None, para_indices: Optional[List[int]] = None, bold: Optional[bool] = None, italic: Optional[bool] = None,
                underline: Optional[bool] = None, strike: Optional[bool] = None, font_name: Optional[str] = None, font_size: Optional[float] = None,
                font_color: Optional[str] = None, highlight: Optional[str] = None, all_caps: Optional[bool] = None, small_caps: Optional[bool] = None,
                superscript: Optional[bool] = None, subscript: Optional[bool] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    count = 0
    paras_to_format = []

    if para_indices:
        for idx in para_indices:
            if 0 <= idx < len(doc.paragraphs):
                paras_to_format.append(doc.paragraphs[idx])
    elif match:
        for para in doc.paragraphs:
            if re.search(match, para.text, re.IGNORECASE):
                paras_to_format.append(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if re.search(match, para.text, re.IGNORECASE):
                            paras_to_format.append(para)
    else:
        return errors.err("formatting", "format_text", "Either --match or --para-indices must be specified.")

    for para in paras_to_format:
        for run in para.runs:
            if match and not re.search(match, run.text, re.IGNORECASE):
                continue
            _apply_run_format(run, bold, italic, underline, strike, font_name,
                              font_size, font_color, highlight, all_caps, small_caps,
                              superscript, subscript)
            count += 1

    if count == 0:
        return errors.warn("formatting", "format_text", "No text found to format.")
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"{count} runs formatted.")


def _apply_run_format(run: Any, bold: Optional[bool], italic: Optional[bool], underline: Optional[bool], strike: Optional[bool], font_name: Optional[str],
                      font_size: Optional[float], font_color: Optional[str], highlight: Optional[str], all_caps: Optional[bool], small_caps: Optional[bool],
                      superscript: Optional[bool], subscript: Optional[bool]) -> None:
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if underline is not None: run.underline = underline
    if strike is not None: run.font.strike = strike
    if font_name: run.font.name = font_name
    if font_size: run.font.size = Pt(font_size)
    if font_color:
        color = font_color.lstrip('#')
        run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    if highlight:
        from docx.enum.text import WD_COLOR_INDEX
        highlight_map = {
            'yellow': WD_COLOR_INDEX.YELLOW, 'green': WD_COLOR_INDEX.BRIGHT_GREEN,
            'cyan': WD_COLOR_INDEX.TURQUOISE, 'pink': WD_COLOR_INDEX.PINK,
            'red': WD_COLOR_INDEX.RED, 'blue': WD_COLOR_INDEX.BLUE,
            'gray': WD_COLOR_INDEX.GRAY_25, 'darkgray': WD_COLOR_INDEX.GRAY_50,
        }
        if highlight.lower() in highlight_map:
            run.font.highlight_color = highlight_map[highlight.lower()]
    if all_caps is not None: run.font.all_caps = all_caps
    if small_caps is not None: run.font.small_caps = small_caps
    if superscript is not None: run.font.superscript = superscript
    if subscript is not None: run.font.subscript = subscript


def format_paragraph(doc_path: str, para_indices: List[int], alignment: Optional[str] = None, line_spacing: Optional[float] = None,
                     space_before: Optional[float] = None, space_after: Optional[float] = None, left_indent: Optional[float] = None,
                     right_indent: Optional[float] = None, first_line_indent: Optional[float] = None, keep_together: Optional[bool] = None,
                     keep_with_next: Optional[bool] = None, page_break_before: Optional[bool] = None,
                     widow_control: Optional[bool] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    count = 0
    for idx in para_indices:
        if 0 <= idx < len(doc.paragraphs):
            pf = doc.paragraphs[idx].paragraph_format
            if alignment and alignment.lower() in ALIGN_MAP:
                pf.alignment = ALIGN_MAP[alignment.lower()]
            if line_spacing is not None:
                pf.line_spacing = Pt(line_spacing) if line_spacing > 3 else line_spacing
            if space_before is not None: pf.space_before = Pt(space_before)
            if space_after is not None: pf.space_after = Pt(space_after)
            if left_indent is not None: pf.left_indent = Cm(left_indent)
            if right_indent is not None: pf.right_indent = Cm(right_indent)
            if first_line_indent is not None: pf.first_line_indent = Cm(first_line_indent)
            if keep_together is not None: pf.keep_together = keep_together
            if keep_with_next is not None: pf.keep_with_next = keep_with_next
            if page_break_before is not None: pf.page_break_before = page_break_before
            if widow_control is not None: pf.widow_control = widow_control
            count += 1

    if count == 0:
        return errors.warn("formatting", "format_paragraph", "No paragraphs found at specified indices.")
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"{count} paragraphs formatted.")


@functools.lru_cache(maxsize=32)
def list_styles(doc_path: str, style_type: Optional[str] = None, json_mode: bool = False) -> Union[str, Dict[str, Any]]:
    doc, err = load_document(doc_path)
    if err:
        return err

    # Use enum comparison instead of fragile str() matching
    type_map = {
        'paragraph': WD_STYLE_TYPE.PARAGRAPH,
        'character': WD_STYLE_TYPE.CHARACTER,
        'table': WD_STYLE_TYPE.TABLE,
        'list': WD_STYLE_TYPE.LIST,
    }

    styles_data = []
    lines = ["=== STYLES ==="]
    for style in doc.styles:
        if style_type and style_type.lower() in type_map:
            if style.type != type_map[style_type.lower()]:
                continue
        entry = {
            "name": style.name,
            "type": str(style.type),
            "builtin": style.builtin,
        }
        styles_data.append(entry)
        lines.append(f"  [{style.type}] {style.name} (built-in: {style.builtin})")

    if json_mode:
        return {"status": "SUCCESS", "styles": styles_data}
    return "\n".join(lines)


def apply_style(doc_path: str, para_indices: List[int], style_name: str, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    count = 0
    for idx in para_indices:
        if 0 <= idx < len(doc.paragraphs):
            try:
                doc.paragraphs[idx].style = style_name
                count += 1
            except KeyError:
                return errors.err("formatting", "apply_style", f"Style '{style_name}' not found in document styles.")
    if count == 0:
        return errors.warn("formatting", "apply_style", "No paragraphs found at specified indices.")
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"{count} paragraphs applied style '{style_name}'.")
