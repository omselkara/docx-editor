"""Smart features: summary map, page-based reading, format painter, template system, language detection."""
import re
import os
from typing import List, Optional, Dict, Any, Union, Set, Callable
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx_engine.core import load_document, save_document, has_page_break
from docx_engine.constants import LINES_PER_PAGE, CHARS_PER_LINE
from docx_engine import errors


# ===================== SUMMARY MAP =====================

def summary_map(doc_path: str) -> str:
    """Create a compact overview of the entire document structure."""
    doc, err = load_document(doc_path)
    if err:
        return err

    body = doc.element.body
    img_count = len(list(body.iter(qn('w:drawing')))) + len(list(body.iter(qn('w:pict'))))
    comment_count = len(list(body.iter(qn('w:commentRangeStart'))))

    lines = ["=== DOCUMENT MAP ==="]
    lines.append(f"Paragraphs: {len(doc.paragraphs)} | Tables: {len(doc.tables)} | "
                 f"Sections: {len(doc.sections)} | Images: {img_count} | Comments: {comment_count}")
    lines.append("")

    # Approximate page tracking
    current_page = 1
    line_count = 0

    last_page_shown = 0
    empty_streak = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        # Check for page breaks
        if has_page_break(para) and line_count > 0:
            current_page += 1
            line_count = 0

        # Skip empty paragraphs — just count them
        if not text:
            empty_streak += 1
            line_count += 1
            if line_count >= LINES_PER_PAGE:
                current_page += 1
                line_count = 0
            continue

        # Flush empty streak as a single compact note
        if empty_streak > 1:
            lines.append(f"  ... ({empty_streak} empty lines)")
        empty_streak = 0

        if current_page != last_page_shown:
            lines.append(f"--- Page ~{current_page} ---")
            last_page_shown = current_page

        # Compact representation
        preview = text[:60] + "..." if len(text) > 60 else text

        marker = ""
        if style.startswith("Heading"):
            level = style.replace("Heading ", "").strip()
            marker = f"H{level} "
        elif "List" in style:
            marker = "• "

        lines.append(f"  [P{i}] {marker}{preview}")

        # Track estimated lines
        estimated_lines = max(1, len(text) // 80 + 1)
        line_count += estimated_lines
        if line_count >= LINES_PER_PAGE:
            current_page += 1
            line_count = 0

    # Add table summary
    if doc.tables:
        lines.append("\n--- Tables ---")
        for t_idx, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            first_cell = table.rows[0].cells[0].text[:30] if rows > 0 else ""
            lines.append(f"  [T{t_idx}] {rows}x{cols} — '{first_cell}...'")

    lines.append(f"\n--- Estimated pages: ~{current_page} ---")
    return "\n".join(lines)


# ===================== PAGE-BASED READING =====================

def read_page(doc_path: str, page_number: int) -> str:
    """Read content from an approximate page number."""
    doc, err = load_document(doc_path)
    if err:
        return err

    current_page = 1
    line_count = 0

    page_content = []
    capturing = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        style = para.style.name

        # Check for page breaks
        if has_page_break(para) and line_count > 0:
            current_page += 1
            line_count = 0
            if capturing:
                break  # We've gone past the target page

        if current_page == page_number:
            capturing = True
            page_content.append(f"[P{i}] ({style}) {text}")
        elif current_page > page_number and capturing:
            break

        estimated_lines = max(1, len(text) // 80 + 1)
        line_count += estimated_lines
        if line_count >= LINES_PER_PAGE:
            current_page += 1
            line_count = 0
            if capturing:
                break

    if not page_content:
        return errors.err("smart_features", "read_page", f"Page {page_number} not found. Document is approximately {current_page} pages.")
    return f"=== PAGE {page_number} (approximate) ===\n" + "\n".join(page_content)


# ===================== FORMAT PAINTER =====================

def clone_format(doc_path: str, source_para: int, target_paras: List[int], output_path: Optional[str] = None) -> str:
    """Clone all formatting from source paragraph to target paragraphs."""
    doc, err = load_document(doc_path)
    if err:
        return err

    paras = doc.paragraphs
    if source_para >= len(paras):
        return errors.err("smart_features", "clone_format", "Invalid source paragraph index.")

    src = paras[source_para]

    # Capture paragraph-level format
    src_pf = src.paragraph_format
    src_style = src.style

    # Capture run-level format from first run
    src_run_fmt = None
    if src.runs:
        r = src.runs[0]
        src_run_fmt = {
            'bold': r.bold,
            'italic': r.italic,
            'underline': r.underline,
            'font_name': r.font.name,
            'font_size': r.font.size,
            'font_color_rgb': r.font.color.rgb if r.font.color and r.font.color.rgb else None,
            'strike': r.font.strike,
            'all_caps': r.font.all_caps,
            'small_caps': r.font.small_caps,
        }

    count = 0
    for idx in target_paras:
        if 0 <= idx < len(paras):
            tgt = paras[idx]
            # Apply style
            tgt.style = src_style

            # Apply paragraph format
            tgt_pf = tgt.paragraph_format
            tgt_pf.alignment = src_pf.alignment
            tgt_pf.line_spacing = src_pf.line_spacing
            tgt_pf.space_before = src_pf.space_before
            tgt_pf.space_after = src_pf.space_after
            tgt_pf.left_indent = src_pf.left_indent
            tgt_pf.right_indent = src_pf.right_indent
            tgt_pf.first_line_indent = src_pf.first_line_indent
            tgt_pf.keep_together = src_pf.keep_together
            tgt_pf.keep_with_next = src_pf.keep_with_next
            tgt_pf.page_break_before = src_pf.page_break_before

            # Apply run format
            if src_run_fmt:
                for run in tgt.runs:
                    if src_run_fmt['bold'] is not None:
                        run.bold = src_run_fmt['bold']
                    if src_run_fmt['italic'] is not None:
                        run.italic = src_run_fmt['italic']
                    if src_run_fmt['underline'] is not None:
                        run.underline = src_run_fmt['underline']
                    if src_run_fmt['font_name']:
                        run.font.name = src_run_fmt['font_name']
                    if src_run_fmt['font_size']:
                        run.font.size = src_run_fmt['font_size']
                    if src_run_fmt['font_color_rgb']:
                        run.font.color.rgb = src_run_fmt['font_color_rgb']
                    if src_run_fmt['strike'] is not None:
                        run.font.strike = src_run_fmt['strike']
                    if src_run_fmt['all_caps'] is not None:
                        run.font.all_caps = src_run_fmt['all_caps']
                    if src_run_fmt['small_caps'] is not None:
                        run.font.small_caps = src_run_fmt['small_caps']
            count += 1

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Format cloned from P{source_para} to {count} paragraphs.")


# ===================== TEMPLATE SYSTEM =====================

def from_template(template_path: str, output_path: str, variables: Optional[str] = None) -> str:
    """Create a new document from a template, replacing {{variables}}."""
    doc, err = load_document(template_path)
    if err:
        return err

    if not variables:
        return errors.err("smart_features", "from_template", "Variables not specified. Format: key1=value1,key2=value2")

    # Parse variables
    var_dict = {}
    for pair in variables.split(","):
        if "=" in pair:
            key, val = pair.split("=", 1)
            var_dict[key.strip()] = val.strip()

    def replace_fn(paragraphs: List[Any]) -> int:
        return _replace_in_paragraph_list(paragraphs, var_dict)

    replaced_count = _execute_on_every_paragraph(doc, replace_fn)

    doc.save(output_path)
    return errors.ok(f"Template applied → {output_path}") + f"\n{replaced_count} variables replaced.\nVariables: {var_dict}"


def list_template_variables(doc_path: str) -> str:
    """Find all {{variable}} placeholders in a document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    variables: Set[str] = set()

    def find_fn(paragraphs: List[Any]) -> int:
        variables.update(_find_variables_in_paragraph_list(paragraphs))
        return 0

    _execute_on_every_paragraph(doc, find_fn)

    if not variables:
        return errors.warn("smart_features", "list_template_variables", "No template variables found.")
    return errors.ok("Found variables: " + ", ".join(sorted(list(variables))))


def _execute_on_every_paragraph(doc: Any, fn: Callable[[List[Any]], int]) -> int:
    """Helper to execute a function on every paragraph list in the document (body, tables, headers/footers)."""
    total = 0
    # Main body
    total += fn(doc.paragraphs)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += fn(cell.paragraphs)

    # Headers/Footers
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                total += fn(hf.paragraphs)
    return total


def _replace_in_paragraph_list(paragraphs: List[Any], var_dict: Dict[str, str]) -> int:
    """Helper to replace variables in a list of paragraphs."""
    count = 0
    for para in paragraphs:
        for key, val in var_dict.items():
            placeholder = "{{" + key + "}}"
            if placeholder in para.text:
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, val)
                        count += 1
    return count


def _find_variables_in_paragraph_list(paragraphs: List[Any]) -> Set[str]:
    """Helper to find {{variable}} placeholders in a list of paragraphs."""
    pattern = r'\{\{(\w+)\}\}'
    variables = set()
    for para in paragraphs:
        matches = re.findall(pattern, para.text)
        variables.update(matches)
    return variables


# ===================== LANGUAGE DETECTION =====================

def detect_language(doc_path: str) -> str:
    """Detect the primary language of the document."""
    doc, err = load_document(doc_path)
    if err:
        return err

    # Collect all text
    all_text = " ".join(p.text for p in doc.paragraphs if p.text.strip())

    if not all_text:
        return errors.warn("smart_features", "detect_language", "Document is empty, language could not be detected.")

    # Simple heuristic-based detection
    char_count = len(all_text)

    # Turkish-specific characters
    tr_chars = set('çÇğĞıİöÖşŞüÜ')
    tr_count = sum(1 for c in all_text if c in tr_chars)

    # German-specific
    de_chars = set('äÄöÖüÜß')
    de_count = sum(1 for c in all_text if c in de_chars)

    # French-specific
    fr_chars = set('àâæçéèêëïîôùûüÿœÀÂÆÇÉÈÊËÏÎÔÙÛÜŸŒ')
    fr_count = sum(1 for c in all_text if c in fr_chars)

    # Spanish-specific
    es_chars = set('áéíóúñüÁÉÍÓÚÑÜ¿¡')
    es_count = sum(1 for c in all_text if c in es_chars)

    # Arabic/Persian
    arabic_count = sum(1 for c in all_text if '\u0600' <= c <= '\u06FF')

    # Chinese/Japanese/Korean
    cjk_count = sum(1 for c in all_text if '\u4e00' <= c <= '\u9fff' or '\u3040' <= c <= '\u30ff')

    # Cyrillic (Russian etc.)
    cyrillic_count = sum(1 for c in all_text if '\u0400' <= c <= '\u04FF')

    # Check XML language settings too
    xml_langs = set()
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:
                lang = rPr.find(qn('w:lang'))
                if lang is not None:
                    for attr in ['val', 'bidi', 'eastAsia']:
                        v = lang.get(qn(f'w:{attr}'))
                        if v:
                            xml_langs.add(v)

    # Determine primary language
    scores = {
        'Turkish (tr)': tr_count * 10,
        'German (de)': de_count * 10,
        'French (fr)': fr_count * 10,
        'Spanish (es)': es_count * 10,
        'Arabic (ar)': arabic_count * 5,
        'CJK (cjk)': cjk_count * 5,
        'Russian/Cyrillic (ru)': cyrillic_count * 5,
        'English (en)': 1,  # Default fallback
    }

    # ASCII-dominant with no special chars likely English
    ascii_ratio = sum(1 for c in all_text if c.isascii()) / max(char_count, 1)
    if ascii_ratio > 0.95 and max(tr_count, de_count, fr_count, es_count) == 0:
        scores['English (en)'] = char_count

    primary = max(scores, key=scores.get)

    lines = ["=== LANGUAGE ANALYSIS ==="]
    lines.append(f"  Primary language: {primary}")
    lines.append(f"  Total characters: {char_count}")
    lines.append(f"  ASCII ratio: {ascii_ratio:.1%}")
    if xml_langs:
        lines.append(f"  XML language tags: {', '.join(xml_langs)}")
    lines.append("\n  Language scores:")
    for lang, score in sorted(scores.items(), key=lambda x: -x[1]):
        if score > 0:
            lines.append(f"    {lang}: {score}")

    return "\n".join(lines)


# ===================== WORD COUNT PER SECTION =====================

def word_count(doc_path: str) -> str:
    """Detailed word count broken down by section/heading."""
    doc, err = load_document(doc_path)
    if err:
        return err

    total_words = 0
    total_chars = 0
    section_counts = []
    current_section = {"name": "(Start)", "words": 0, "chars": 0, "paras": 0}

    for para in doc.paragraphs:
        text = para.text.strip()
        words = len(text.split()) if text else 0
        chars = len(text)

        if para.style.name.startswith('Heading'):
            if current_section["paras"] > 0:
                section_counts.append(current_section)
            current_section = {"name": text or "(Untitled)", "words": 0, "chars": 0, "paras": 0}

        current_section["words"] += words
        current_section["chars"] += chars
        current_section["paras"] += 1
        total_words += words
        total_chars += chars

    if current_section["paras"] > 0:
        section_counts.append(current_section)

    # Table word count
    table_words = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_words += len(cell.text.split())

    lines = ["=== WORD COUNT ==="]
    lines.append(f"  Total words: {total_words}")
    lines.append(f"  Total characters: {total_chars}")
    lines.append(f"  Table word count: {table_words}")
    lines.append(f"  Grand total: {total_words + table_words}")
    lines.append(f"\n  Section based:")
    for sec in section_counts:
        lines.append(f"    [{sec['paras']} paras] {sec['name']}: {sec['words']} words, {sec['chars']} characters")

    return "\n".join(lines)
