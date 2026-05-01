"""Advanced features: images, headers/footers, page layout, lists, hyperlinks, page breaks, TOC."""
import os
from typing import Any, Optional

from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt

from docx_engine import errors
from docx_engine.core import load_document, save_document

# ===================== IMAGES =====================

def insert_image(doc_path: str, image_path: str, width: Optional[float] = None, height: Optional[float] = None, after_para: Optional[int] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if not os.path.exists(image_path):
        return errors.err("advanced", "insert_image", f"Image file not found: {image_path}")

    # Available content area in EMU (page minus margins)
    section = doc.sections[0]
    avail_w = section.page_width - section.left_margin - section.right_margin
    avail_h = section.page_height - section.top_margin - section.bottom_margin

    # Get real image dimensions in EMU via PIL
    img_w_emu = img_h_emu = None
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(image_path) as _img:
            px_w, px_h = _img.size
            dpi = _img.info.get("dpi", (96, 96))
            dpi_x = float(dpi[0]) if isinstance(dpi, (tuple, list)) else 96.0
            dpi_y = float(dpi[1]) if isinstance(dpi, (tuple, list)) else 96.0
            from docx_engine.constants import EMU_PER_INCH as _EMU
            img_w_emu = int(px_w / dpi_x * _EMU)
            img_h_emu = int(px_h / dpi_y * _EMU)
    except Exception:
        pass

    if width or height:
        # User-specified dimensions: convert then clamp to page bounds
        req_w = Inches(width) if width else None
        req_h = Inches(height) if height else None

        if req_w and req_h:
            # Both given: scale down proportionally if either exceeds available area
            scale = min(avail_w / req_w, avail_h / req_h, 1.0)
            kwargs = {'width': int(req_w * scale), 'height': int(req_h * scale)}
        elif req_w:
            kwargs = {'width': min(req_w, avail_w)}
        else:
            kwargs = {'height': min(req_h, avail_h)}
    else:
        # No dimensions given: fit entire image within available area, never upscale
        if img_w_emu and img_h_emu:
            scale = min(avail_w / img_w_emu, avail_h / img_h_emu, 1.0)
            kwargs = {'width': int(img_w_emu * scale), 'height': int(img_h_emu * scale)}
        else:
            # PIL unavailable: at least prevent horizontal overflow
            kwargs = {'width': avail_w}

    if after_para is not None and 0 <= after_para < len(doc.paragraphs):
        new_p = doc.add_paragraph()
        run = new_p.add_run()
        run.add_picture(image_path, **kwargs)
        doc.paragraphs[after_para]._p.addnext(new_p._p)
    else:
        doc.add_picture(image_path, **kwargs)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Image added: {image_path}")


def list_images(doc_path: str) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    images = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            images.append(f"  [IMG{i}] {rel.target_ref} (type: {rel.reltype.split('/')[-1]})")

    if not images:
        return errors.err("advanced", "list_images", "No images found in document.")
    return "=== IMAGES ===\n" + "\n".join(images)


# ===================== HEADERS & FOOTERS =====================

def set_header(doc_path: str, text: str, section_index: int = 0, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if section_index >= len(doc.sections):
        return errors.err("advanced", "set_header", "Invalid section index.")

    section = doc.sections[section_index]
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        header.paragraphs[0].text = text
    else:
        header.add_paragraph(text)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Header set: '{text}'")


def set_footer(doc_path: str, text: str, add_page_number: bool = False, section_index: int = 0, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    if section_index >= len(doc.sections):
        return errors.err("advanced", "set_footer", "Invalid section index.")

    section = doc.sections[section_index]
    footer = section.footer
    footer.is_linked_to_previous = False

    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()

    para.text = text

    if add_page_number:
        run = para.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = para.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = para.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok("Footer set.")


def read_header(doc_path: str, section_index: int = 0) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    section = doc.sections[section_index]
    header = section.header
    text = "\n".join(p.text for p in header.paragraphs)
    return f"=== HEADER (Section {section_index}) ===\n{text if text.strip() else '(empty)'}"


def read_footer(doc_path: str, section_index: int = 0) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    section = doc.sections[section_index]
    footer = section.footer
    text = "\n".join(p.text for p in footer.paragraphs)
    return f"=== FOOTER (Section {section_index}) ===\n{text if text.strip() else '(empty)'}"


# ===================== PAGE LAYOUT =====================

def set_margins(doc_path: str, top: Optional[float] = None, bottom: Optional[float] = None, left: Optional[float] = None, right: Optional[float] = None,
                section_index: Optional[int] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    sections = [doc.sections[section_index]] if section_index is not None else doc.sections
    for section in sections:
        if top is not None: section.top_margin = Cm(top)
        if bottom is not None: section.bottom_margin = Cm(bottom)
        if left is not None: section.left_margin = Cm(left)
        if right is not None: section.right_margin = Cm(right)
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok("Margins set.")


def set_orientation(doc_path: str, orientation: str, section_index: Optional[int] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err
    sections = [doc.sections[section_index]] if section_index is not None else doc.sections
    for section in sections:
        if orientation.lower() == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            new_width, new_height = section.page_height, section.page_width
            section.page_width = min(new_width, new_height)
            section.page_height = max(new_width, new_height)
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Orientation set to {orientation}")


def set_page_size(doc_path: str, preset: Optional[str] = None, width: Optional[float] = None, height: Optional[float] = None,
                  section_index: Optional[int] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    presets = {
        'a4': (21.0, 29.7), 'letter': (21.59, 27.94), 'legal': (21.59, 35.56),
        'a3': (29.7, 42.0), 'a5': (14.8, 21.0), 'b5': (17.6, 25.0),
    }

    if preset and preset.lower() in presets:
        width, height = presets[preset.lower()]

    if width is None or height is None:
        return errors.err("advanced", "set_page_size", "Specify page size (preset or width/height in cm).")

    sections = [doc.sections[section_index]] if section_index is not None else doc.sections
    for section in sections:
        section.page_width = Cm(width)
        section.page_height = Cm(height)
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Page size set to {width}x{height}cm")


def insert_page_break(doc_path: str, after_para: Optional[int] = None, output_path: Optional[str] = None) -> str:
    from docx.enum.text import WD_BREAK
    doc, err = load_document(doc_path)
    if err:
        return err
    if after_para is not None and 0 <= after_para < len(doc.paragraphs):
        run = doc.paragraphs[after_para].add_run()
        run.add_break(WD_BREAK.PAGE)
    else:
        doc.add_page_break()
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok("Page break inserted.")


def insert_section_break(doc_path: str, break_type: str = 'new_page', after_para: Optional[int] = None, output_path: Optional[str] = None) -> str:
    from docx.enum.section import WD_SECTION_START
    doc, err = load_document(doc_path)
    if err:
        return err

    type_map = {
        'new_page': WD_SECTION_START.NEW_PAGE,
        'continuous': WD_SECTION_START.CONTINUOUS,
        'even_page': WD_SECTION_START.EVEN_PAGE,
        'odd_page': WD_SECTION_START.ODD_PAGE,
    }

    doc.add_section(type_map.get(break_type, WD_SECTION_START.NEW_PAGE))
    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Section break inserted: {break_type}")


# ===================== LISTS =====================

def insert_list(doc_path: str, items: str, list_type: str = 'bullet', after_para: Optional[int] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    style = 'List Bullet' if list_type == 'bullet' else 'List Number'
    item_list = [i.strip() for i in items.split(";")]

    # Check if the style exists in the document; fall back gracefully if not
    style_exists = any(s.name == style for s in doc.styles)
    fallback_used = False

    if style_exists:
        for item in item_list:
            doc.add_paragraph(item, style=style)
    else:
        # Fallback: add plain paragraphs with marker characters
        fallback_used = True
        for idx, item in enumerate(item_list):
            prefix = "•  " if list_type == 'bullet' else f"{idx+1}.  "
            doc.add_paragraph(prefix + item)

    result = save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"{list_type} list with {len(item_list)} items added.")
    if fallback_used:
        result += "\n" + errors.warn("advanced", "insert_list", f"Style '{style}' not found in document. Used plain text markers as fallback. Run apply_style to fix.")
    return result


# ===================== HYPERLINKS =====================

def insert_hyperlink(doc_path: str, text: str, url: str, after_para: Optional[int] = None, output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    para = doc.add_paragraph()
    _add_hyperlink(para, url, text)

    if after_para is not None and 0 <= after_para < len(doc.paragraphs):
        doc.paragraphs[after_para]._p.addnext(para._p)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok(f"Hyperlink added: {text} -> {url}")


def _add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
    new_run = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr><w:t>{text}</w:t></w:r>')
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ===================== TABLE OF CONTENTS =====================

def insert_toc(doc_path: str, title: str = "Table of Contents", output_path: Optional[str] = None) -> str:
    doc, err = load_document(doc_path)
    if err:
        return err

    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    # TOC field
    fld_begin = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin"/></w:r>')
    fld_instr = parse_xml(f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>')
    fld_sep = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="separate"/></w:r>')
    fld_text = parse_xml(f'<w:r {nsdecls("w")}><w:t>[Right-click and select "Update Field" in Word to update the Table of Contents]</w:t></w:r>')
    fld_end = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>')

    toc_para = doc.add_paragraph()
    toc_para._p.append(fld_begin)
    toc_para._p.append(fld_instr)
    toc_para._p.append(fld_sep)
    toc_para._p.append(fld_text)
    toc_para._p.append(fld_end)

    return save_document(doc, doc_path, output_path) + "\n" + errors.ok("Table of Contents field added.")
